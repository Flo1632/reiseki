#!/usr/bin/env python3
"""
Local File Agent — Ollama + FastAPI
====================================
Start:  python agent.py
Config: AGENT_MODEL=qwen2.5-coder:7b  AGENT_ROOT=/dein/pfad  python agent.py
"""

import asyncio
import base64
import datetime
import logging
from difflib import SequenceMatcher
import io
import json
import os
import queue as _queue
import re
import socket
import sqlite3
import threading
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger("agent")

from fastapi import FastAPI, File, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import ollama

# ── Config ───────────────────────────────────────────────────────────────────
MODEL                  = os.environ.get("AGENT_MODEL", "qwen2.5-coder:7b")
ORIGINAL_ROOT          = Path(os.environ.get("AGENT_ROOT", ".")).resolve()
ROOT                   = ORIGINAL_ROOT
HOST                   = os.environ.get("AGENT_HOST", "127.0.0.1")
DB_PATH                = Path(__file__).parent / "agent_memory.db"
CONTEXT_COMPRESS_AFTER = 4   # compress history after this many tool calls per turn
HISTORY_MAX_MESSAGES   = 20  # max messages kept across turns before compression
MEMORY_TOP_K           = 8   # relevant memories injected per request
TOOLS_TOP_K            = 4   # dynamic tools selected per request (+ always-on)
SEARCH_CONFIRM_TIMEOUT = 30  # seconds to wait for user confirmation before auto-cancel

# ── Web-search confirmation state (single-user) ───────────────────────────────
_search_confirm_event    = threading.Event()
_search_confirm_approved = False

# ── Concurrency lock (single-user app — serialise requests to protect globals) ─
_request_lock = asyncio.Lock()

def _local_ip() -> str:
    """Return the local LAN IP without making any outbound connection."""
    try:
        hostname = socket.gethostname()
        candidates = socket.getaddrinfo(hostname, None, socket.AF_INET)
        for item in candidates:
            ip = item[4][0]
            if not ip.startswith("127."):
                return ip
    except Exception:
        pass
    return "127.0.0.1"

# ── Persistent memory (SQLite) ────────────────────────────────────────────────
def _init_db() -> None:
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("""CREATE TABLE IF NOT EXISTS config (
            key   TEXT PRIMARY KEY,
            value TEXT NOT NULL
        )""")
        conn.execute("""CREATE TABLE IF NOT EXISTS memories (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT    DEFAULT (datetime('now')),
            content    TEXT    NOT NULL,
            category   TEXT    DEFAULT 'general',
            importance REAL    DEFAULT 0.5
        )""")
        # Migrate existing DB: add importance column if missing
        try:
            conn.execute("ALTER TABLE memories ADD COLUMN importance REAL DEFAULT 0.5")
        except sqlite3.OperationalError:
            pass  # column already exists
        conn.execute("""CREATE TABLE IF NOT EXISTS appointments (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            due_at      TEXT    NOT NULL,
            title       TEXT    NOT NULL,
            description TEXT    DEFAULT '',
            notified    INTEGER DEFAULT 0
        )""")
        conn.execute("""CREATE TABLE IF NOT EXISTS chat_log (
            id         INTEGER PRIMARY KEY AUTOINCREMENT,
            created_at TEXT    DEFAULT (datetime('now')),
            role       TEXT    NOT NULL,
            content    TEXT    NOT NULL
        )""")

_init_db()

def _warmup_model() -> None:
    """Load the model into Ollama's memory at startup so the first user request is fast."""
    try:
        ollama.chat(model=MODEL, messages=[{"role": "user", "content": "hi"}])
        print(f"✅  Model '{MODEL}' warmed up and ready.")
    except Exception as e:
        print(f"⚠️  Model warmup failed: {e}")

threading.Thread(target=_warmup_model, daemon=True).start()

def _log_message(role: str, content: str) -> None:
    """Persist a single chat turn to the chat_log table."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "INSERT INTO chat_log (role, content) VALUES (?,?)",
            (role, content[:10000])
        )

def _cfg_get(key: str, default: str = "") -> str:
    with sqlite3.connect(DB_PATH) as conn:
        row = conn.execute("SELECT value FROM config WHERE key=?", (key,)).fetchone()
    return row[0] if row else default

def _cfg_set(key: str, value: str) -> None:
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("INSERT OR REPLACE INTO config (key, value) VALUES (?,?)", (key, value))

def _relevant_memories(query: str, top_k: int = MEMORY_TOP_K) -> list[tuple]:
    """Return top-K memories scored by relevance to query (overlap + importance + recency)."""
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            "SELECT created_at, category, content, importance FROM memories ORDER BY id DESC LIMIT 100"
        ).fetchall()
    if not rows:
        return []
    now = datetime.datetime.now()
    scored = []
    for created_at, category, content, importance in rows:
        overlap  = SequenceMatcher(None, query.lower(), content.lower()).ratio()
        try:
            created  = datetime.datetime.fromisoformat(created_at)
            hours    = (now - created).total_seconds() / 3600
            recency  = 1.0 / (1.0 + hours)
        except Exception:
            recency = 0.5
        score = overlap * 0.5 + (importance or 0.5) * 0.3 + recency * 0.2
        scored.append((score, created_at, category, content))
    scored.sort(reverse=True)
    return scored[:top_k]

def _build_system(query: str = "") -> str:
    name = _cfg_get("agent_name", "File Agent")
    goal = _cfg_get("user_goal", "")
    now  = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    base = (
        f"You are {name}, a local file and productivity agent.\n"
        f"Working root directory: {ROOT}\n"
        f"Current date/time: {now}\n"
        f"Always use relative paths from the root. Be concise and helpful.\n"
        f"Only use tools when strictly necessary. For simple questions or conversational replies, answer directly without calling any tool.\n"
        f"Use the minimum number of tool calls needed — do not chain tools unnecessarily.\n"
        f"IMPORTANT — Tool chaining rule: When creating a file or document (write_file, create_docx, create_xlsx), "
        f"you MUST use the ACTUAL content returned by previous tool calls (e.g. web_search results, read_file content). "
        f"Never write placeholder text or invent content — copy the real tool result into the document.\n"
        f"Once you have enough information to answer and the tool calls have satisfied the requirement, stop calling tools and respond directly."
    )
    # Root directory listing so the model knows what files actually exist
    try:
        root_entries = sorted(ROOT.iterdir(), key=lambda p: (p.is_dir(), p.name))
        root_listing = "\n".join(
            ("📁 " if e.is_dir() else "📄 ") + e.name for e in root_entries[:50]
        )
        base += f"\n\nFiles and folders in the working root:\n{root_listing}"
    except Exception:
        pass
    if goal:
        base += f"\n\nUser's goal: {goal}"
    # Relevance-scored memories (not FIFO) — top-K most relevant to current query
    mem_rows = _relevant_memories(query or goal)
    if mem_rows:
        mem_lines = "\n".join(f"[{r[1]}] ({r[2]}) {r[3]}" for r in mem_rows)
        base += f"\n\nRelevant memories — use these to personalise your responses:\n{mem_lines}"
    return base

# ── Safety helper ─────────────────────────────────────────────────────────────
def _safe(path: str) -> Path | None:
    target = (ROOT / path).resolve()
    return target if target == ROOT or ROOT in target.parents else None

# ── Tool implementations ──────────────────────────────────────────────────────
def list_directory(path: str = ".") -> str:
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    try:
        entries = sorted(t.iterdir(), key=lambda e: (e.is_file(), e.name.lower()))
        lines = [f"{'📁' if e.is_dir() else '📄'} {e.name}" for e in entries]
        return "\n".join(lines) if lines else "(empty)"
    except Exception as e:
        logger.error("list_directory(%s): %s", path, e)
        return "Error: Could not list directory"

def read_file(path: str) -> str:
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    suffix = t.suffix.lower()
    try:
        if suffix == ".docx":
            from docx import Document
            doc = Document(t)
            content = "\n".join(p.text for p in doc.paragraphs)
        elif suffix in (".xlsx", ".xls"):
            import pandas as pd
            df = pd.read_excel(t)
            content = df.to_string(index=False)
        else:
            content = t.read_text(encoding="utf-8")
        return content[:8000] + "\n[... truncated]" if len(content) > 8000 else content
    except Exception as e:
        logger.error("read_file(%s): %s", path, e)
        return "Error: Could not read file"

_PROTECTED_FILES = {
    "agent.py", "launcher.py", "CLAUDE.md", "agent_memory.db",
    "requirements.txt", "install.sh", "install.bat",
}

def write_file(path: str, content: str) -> str:
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    if t.name in _PROTECTED_FILES:
        return f"Error: '{t.name}' is a protected system file and cannot be overwritten"
    if t.suffix.lower() in (".docx", ".xlsx", ".xls"):
        return f"Error: Use create_docx or create_xlsx for .docx/.xlsx files — write_file only writes plain text."
    if t.exists():
        return f"Error: File '{path}' already exists. Choose a different filename or ask the user whether to overwrite."
    try:
        t.parent.mkdir(parents=True, exist_ok=True)
        t.write_text(content, encoding="utf-8")
        return f"✅ {len(content)} bytes written → {path}"
    except Exception as e:
        logger.error("write_file(%s): %s", path, e)
        return "Error: Could not write file"

def create_directory(path: str) -> str:
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    try:
        t.mkdir(parents=True, exist_ok=True)
        return f"✅ Directory created: {path}"
    except Exception as e:
        logger.error("create_directory(%s): %s", path, e)
        return "Error: Could not create directory"

# Domains blocked in web search results (malware, tracking, spam)
_BLOCKED_DOMAINS = {
    "malware.com", "phishing.com",          # placeholder examples
    "doubleclick.net", "googleadservices.com",
    "tracking.com", "clickbait.info",
}

def _is_safe_url(url: str) -> bool:
    """Return True only for HTTPS URLs whose host is not in the blocklist."""
    try:
        from urllib.parse import urlparse
        p = urlparse(url)
        if p.scheme != "https":
            return False
        host = p.hostname or ""
        # Strip leading 'www.'
        host = host.removeprefix("www.")
        return host not in _BLOCKED_DOMAINS
    except Exception:
        return False

def web_search(query: str, max_results: int = 5) -> str:
    """Search the web via DuckDuckGo. Only the query string is sent — no user data."""
    # Cap query length to prevent data exfiltration via search queries
    if len(query) > 200:
        return "Error: Query too long (max. 200 characters)"
    max_results = min(max(1, max_results), 10)
    try:
        from ddgs import DDGS
        # safesearch="on" filters adult/harmful content at the DuckDuckGo level
        with DDGS() as ddgs:
            raw = list(ddgs.text(query, max_results=max_results * 2, safesearch="on"))
        # Keep only HTTPS results not on the blocklist
        results = [r for r in raw if _is_safe_url(r.get("href", ""))][:max_results]
        if not results:
            return "No safe results found."
        lines = [f"**{r['title']}**\n{r['href']}\n{r['body']}" for r in results]
        return "\n\n---\n\n".join(lines)
    except Exception as e:
        return f"Search error: {e}"

def save_memory(content: str, category: str = "general", importance: float = 0.5) -> str:
    importance = max(0.0, min(1.0, float(importance)))
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "INSERT INTO memories (content, category, importance) VALUES (?,?,?)",
            (content, category, importance)
        )
    return f"✅ Memory saved (category: {category}, importance: {importance})"

def list_memories(category: str = "") -> str:
    with sqlite3.connect(DB_PATH) as conn:
        if category:
            rows = conn.execute(
                "SELECT created_at, category, content FROM memories WHERE category=? ORDER BY id DESC LIMIT 50",
                (category,)
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT created_at, category, content FROM memories ORDER BY id DESC LIMIT 50"
            ).fetchall()
    if not rows:
        return "No memories found."
    return "\n\n".join(f"[{r[0]}] ({r[1]})\n{r[2]}" for r in rows)

def add_appointment(due_at: str, title: str, description: str = "") -> str:
    """due_at must be ISO format: YYYY-MM-DD HH:MM"""
    try:
        datetime.datetime.fromisoformat(due_at)
    except ValueError:
        return "Error: due_at must be in format YYYY-MM-DD HH:MM"
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "INSERT INTO appointments (due_at, title, description) VALUES (?,?,?)",
            (due_at, title, description)
        )
    return f"✅ Appointment saved: '{title}' on {due_at}"

def list_appointments() -> str:
    now = datetime.datetime.now().isoformat(sep=" ", timespec="minutes")
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            "SELECT due_at, title, description FROM appointments ORDER BY due_at LIMIT 50"
        ).fetchall()
    if not rows:
        return "No appointments found."
    lines = []
    for r in rows:
        icon = "✅" if r[0] < now else "📅"
        lines.append(f"{icon} {r[0]} — {r[1]}" + (f"\n   {r[2]}" if r[2] else ""))
    return "\n".join(lines)

def create_docx(path: str, content: str) -> str:
    """Create a Word document (.docx) with the given text content."""
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    try:
        from docx import Document
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        t.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(t))
        return f"✅ Word document created: {path}"
    except Exception as e:
        logger.error("create_docx(%s): %s", path, e)
        return "Error: Could not create Word document"

def create_xlsx(path: str, rows_json: str, **kwargs) -> str:
    """Create an Excel file. rows_json is either:
    - array of arrays: [['Name','Age'],['Alice',30]]  (first row = headers)
    - array of dicts:  [{'Name':'Alice','Age':30}]    (keys become headers automatically)
    """
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    try:
        import openpyxl
        data = json.loads(rows_json)
        if not isinstance(data, list) or not data:
            return "Error: rows_json must be a non-empty JSON array"
        # Normalise: convert array-of-dicts → array-of-arrays
        if isinstance(data[0], dict):
            headers = list(data[0].keys())
            rows = [headers] + [[row.get(h, "") for h in headers] for row in data]
        else:
            rows = data
        wb = openpyxl.Workbook()
        ws = wb.active
        for row in rows:
            ws.append(row if isinstance(row, list) else list(row))
        t.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(t))
        return f"✅ Excel file created: {path} ({len(rows)} rows)"
    except json.JSONDecodeError:
        return "Error: rows_json is not valid JSON"
    except Exception as e:
        logger.error("create_xlsx(%s): %s", path, e)
        return "Error: Could not create Excel file"

def create_csv(path: str, rows_json: str, **kwargs) -> str:
    """Create a CSV file. rows_json is either:
    - array of arrays: [['Name','Age'],['Alice',30]]  (first row = headers)
    - array of dicts:  [{'Name':'Alice','Age':30}]    (keys become headers automatically)
    """
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    try:
        import pandas as pd
        data = json.loads(rows_json)
        if not isinstance(data, list) or not data:
            return "Error: rows_json must be a non-empty JSON array"
        if isinstance(data[0], dict):
            df = pd.DataFrame(data)
        else:
            df = pd.DataFrame(data[1:], columns=data[0])
        t.parent.mkdir(parents=True, exist_ok=True)
        df.to_csv(str(t), index=False)
        return f"✅ CSV file created: {path} ({len(df)} rows, {len(df.columns)} columns)"
    except json.JSONDecodeError:
        return "Error: rows_json is not valid JSON"
    except Exception as e:
        logger.error("create_csv(%s): %s", path, e)
        return "Error: Could not create CSV file"

def csv_to_excel(csv_path: str, excel_path: str = "") -> str:
    """Convert a CSV file to Excel (.xlsx) using pandas."""
    src = _safe(csv_path)
    if not src:
        return "Error: Source path is outside the root directory"
    if not excel_path:
        excel_path = str(Path(csv_path).with_suffix(".xlsx"))
    dst = _safe(excel_path)
    if not dst:
        return "Error: Target path is outside the root directory"
    try:
        import pandas as pd
        df = _read_df(src)
        dst.parent.mkdir(parents=True, exist_ok=True)
        df.to_excel(str(dst), index=False, engine="openpyxl")
        return f"✅ Converted: {csv_path} → {excel_path} ({len(df)} rows, {len(df.columns)} columns)"
    except Exception as e:
        logger.error("csv_to_excel(%s): %s", csv_path, e)
        return "Error: Could not convert CSV to Excel"

def _read_df(t):
    """Read a CSV or Excel file into a DataFrame with auto-detected separator and cleaned column names."""
    import pandas as pd
    ext = t.suffix.lower()
    if ext == ".csv":
        # sep=None + engine='python' auto-detects comma, semicolon, tab, etc.
        df = pd.read_csv(str(t), sep=None, engine="python", encoding_errors="replace")
    elif ext == ".xlsx":
        df = pd.read_excel(str(t), engine="openpyxl")
    elif ext == ".xls":
        df = pd.read_excel(str(t), engine="xlrd")
    else:
        raise ValueError(f"Unsupported file type: {ext}")
    # Strip whitespace and BOM from column names
    df.columns = [str(c).strip().lstrip("\ufeff") for c in df.columns]
    return df

def analyse_data(path: str, **kwargs) -> str:
    """Analyse a CSV or Excel file with pandas and return statistics."""
    t = _safe(path)
    if not t:
        return "Error: Path is outside the root directory"
    try:
        import pandas as pd
        ext = t.suffix.lower()
        if ext not in (".csv", ".xlsx", ".xls"):
            return "Error: Only CSV and Excel files (.csv, .xlsx, .xls) are supported"
        df = _read_df(t)
        lines = [
            f"Rows: {len(df)}, Columns: {len(df.columns)}",
            f"Column names: {', '.join(df.columns.tolist())}",
            "",
            df.describe(include="all").to_string(),
        ]
        return "\n".join(lines)
    except Exception as e:
        logger.error("analyse_data: %s", e)
        return "Error: Could not analyse data"

def create_ascii_art(description: str) -> str:
    """Generate ASCII art for the given description using a dedicated model call without tools."""
    try:
        resp = ollama.chat(
            model=MODEL,
            messages=[{
                "role": "user",
                "content": (
                    f"Create ASCII art for: {description}\n"
                    "Rules: use only printable ASCII characters, no explanations, "
                    "just the ASCII art itself, keep it under 20 lines."
                )
            }],
            options={"num_predict": 400},
        )
        return resp.message.content or "Could not generate ASCII art"
    except Exception as e:
        logger.error("create_ascii_art: %s", e)
        return "Error: Could not generate ASCII art"


def create_chart(path: str, chart_type: str, x_col: str, y_col: str,
                 title: str = "", output_path: str = "chart.png", **kwargs) -> str:
    """Create a chart from a CSV or Excel file and save it as PNG.
    chart_type: 'line', 'bar', 'scatter', 'pie', 'hist'
    """
    t = _safe(path)
    out = _safe(output_path)
    if not t:
        return "Error: Source file is outside the root directory"
    if not out:
        return "Error: Output path is outside the root directory"
    try:
        import pandas as pd
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        df = _read_df(t)
        fig, ax = plt.subplots(figsize=(8, 5))
        ct = chart_type.lower()
        if ct == "line":
            df.plot(x=x_col, y=y_col, ax=ax, kind="line")
        elif ct == "bar":
            df.plot(x=x_col, y=y_col, ax=ax, kind="bar")
        elif ct == "scatter":
            df.plot(x=x_col, y=y_col, ax=ax, kind="scatter")
        elif ct == "pie":
            df.set_index(x_col)[y_col].plot(ax=ax, kind="pie", autopct="%1.1f%%")
        elif ct == "hist":
            df[y_col].plot(ax=ax, kind="hist", bins=20)
        elif ct == "boxplot":
            df.boxplot(column=y_col, by=x_col, ax=ax)
            plt.suptitle("")
        else:
            plt.close(fig)
            return f"Error: Unknown chart type '{chart_type}'. Allowed: line, bar, scatter, pie, hist, boxplot"
        if title:
            ax.set_title(title)
        plt.tight_layout()
        out.parent.mkdir(parents=True, exist_ok=True)
        fig.savefig(str(out), dpi=120)
        plt.close(fig)
        return f"✅ Chart saved: [IMG:{out.name}]"
    except Exception as e:
        logger.error("create_chart(%s): %s", path, e)
        return "Error: Could not create chart"

# ── Dynamic tool selection ────────────────────────────────────────────────────
# always=True  → always included regardless of query
# when_to_use  → free-text scored against the user query via SequenceMatcher
_TOOL_META: dict[str, dict] = {
    "list_directory":   {"always": True,  "when_to_use": "list files find directory explore folder structure what files exist"},
    "read_file":        {"always": True,  "when_to_use": "read file content inspect document open look at text"},
    "write_file":       {"always": True,  "when_to_use": "create write save update file text content output"},
    "create_directory": {"always": False, "when_to_use": "create folder directory organize mkdir new folder"},
    "web_search":       {"always": False, "when_to_use": "search internet web look up online information news current events"},
    "save_memory":      {"always": False, "when_to_use": "remember save memory fact preference goal important information"},
    "list_memories":    {"always": False, "when_to_use": "recall memory remember what do you know history stored"},
    "add_appointment":  {"always": False, "when_to_use": "schedule appointment meeting reminder event deadline calendar"},
    "list_appointments":{"always": False, "when_to_use": "show appointments schedule calendar upcoming events deadlines"},
    "create_docx":      {"always": False, "when_to_use": "create word document docx report letter write formatted"},
    "create_xlsx":      {"always": False, "when_to_use": "create excel spreadsheet xlsx table structured data rows columns"},
    "create_csv":       {"always": False, "when_to_use": "create csv file table rows columns data export comma separated"},
    "csv_to_excel":     {"always": False, "when_to_use": "convert csv to excel xlsx spreadsheet export"},
    "analyse_data":     {"always": False, "when_to_use": "analyze analyse statistics csv excel spreadsheet dataframe numeric columns rows describe"},
    "create_chart":     {"always": False, "when_to_use": "chart graph plot visualization diagram bar line scatter pie histogram"},
    "create_ascii_art": {"always": False, "when_to_use": "ascii art draw design paint sketch picture symbol text art creative visual"},
}

def _select_tools(query: str, top_k: int = TOOLS_TOP_K) -> list:
    """Return always-on tools + top_k most relevant optional tools for this query."""
    always   = [t for t in TOOLS if _TOOL_META.get(t["function"]["name"], {}).get("always")]
    optional = [t for t in TOOLS if not _TOOL_META.get(t["function"]["name"], {}).get("always")]

    def score(tool_def: dict) -> float:
        name = tool_def["function"]["name"]
        when = _TOOL_META.get(name, {}).get("when_to_use", "")
        desc = tool_def["function"].get("description", "")
        text = f"{name} {desc} {when}".lower()
        return SequenceMatcher(None, query.lower(), text).ratio()

    ranked = sorted(optional, key=score, reverse=True)[:top_k]
    return always + ranked

TOOL_MAP = {
    "list_directory":   list_directory,
    "read_file":        read_file,
    "write_file":       write_file,
    "create_directory": create_directory,
    "web_search":       web_search,
    "save_memory":      save_memory,
    "list_memories":    list_memories,
    "add_appointment":  add_appointment,
    "list_appointments": list_appointments,
    "create_docx":      create_docx,
    "create_xlsx":      create_xlsx,
    "create_csv":       create_csv,
    "csv_to_excel":     csv_to_excel,
    "analyse_data":     analyse_data,
    "create_chart":     create_chart,
    "create_ascii_art": create_ascii_art,
}

TOOLS = [
    {"type": "function", "function": {
        "name": "list_directory",
        "description": "Lists files and folders in a directory",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Relative path (default: '.' = root)"}
        }}
    }},
    {"type": "function", "function": {
        "name": "read_file",
        "description": "Reads the content of a file. Supports plain text, .docx and .xlsx/.xls.",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Relative path to the file"}
        }, "required": ["path"]}
    }},
    {"type": "function", "function": {
        "name": "write_file",
        "description": "Creates a new plain text file (.txt, .md, .py, .json etc.). Do NOT use for .docx or .xlsx — use create_docx and create_xlsx instead. Does not overwrite existing files.",
        "parameters": {"type": "object", "properties": {
            "path":    {"type": "string", "description": "Relative path to the file"},
            "content": {"type": "string", "description": "File content"}
        }, "required": ["path", "content"]}
    }},
    {"type": "function", "function": {
        "name": "create_directory",
        "description": "Creates a new directory inside the root directory",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Relative path of the new directory"}
        }, "required": ["path"]}
    }},
    {"type": "function", "function": {
        "name": "web_search",
        "description": "Searches the internet via DuckDuckGo. Only the query is transmitted — no user data.",
        "parameters": {"type": "object", "properties": {
            "query":       {"type": "string",  "description": "Search query"},
            "max_results": {"type": "integer", "description": "Maximum number of results (default: 5)"}
        }, "required": ["query"]}
    }},
    {"type": "function", "function": {
        "name": "save_memory",
        "description": "Saves a memory permanently",
        "parameters": {"type": "object", "properties": {
            "content":    {"type": "string", "description": "Memory content"},
            "category":   {"type": "string", "description": "Category (e.g. 'work', 'personal', 'goal', 'fact', 'result')"},
            "importance": {"type": "number", "description": "Importance 0.0–1.0 (default 0.5). Use 0.9–1.0 for goals/preferences, 0.3–0.5 for general facts."}
        }, "required": ["content"]}
    }},
    {"type": "function", "function": {
        "name": "list_memories",
        "description": "Lists stored memories",
        "parameters": {"type": "object", "properties": {
            "category": {"type": "string", "description": "Filter by category (empty = all)"}
        }}
    }},
    {"type": "function", "function": {
        "name": "add_appointment",
        "description": "Saves an appointment with date and time",
        "parameters": {"type": "object", "properties": {
            "due_at":      {"type": "string", "description": "Date and time in format YYYY-MM-DD HH:MM"},
            "title":       {"type": "string", "description": "Appointment title"},
            "description": {"type": "string", "description": "Optional description"}
        }, "required": ["due_at", "title"]}
    }},
    {"type": "function", "function": {
        "name": "list_appointments",
        "description": "Lists all appointments (past and upcoming)",
        "parameters": {"type": "object", "properties": {}}
    }},
    {"type": "function", "function": {
        "name": "create_docx",
        "description": "Creates a Word document (.docx) with the given text content. Always use this for .docx files — never write_file.",
        "parameters": {"type": "object", "properties": {
            "path":    {"type": "string", "description": "Relative path to the .docx file"},
            "content": {"type": "string", "description": "Text content of the document (line breaks are treated as paragraphs)"}
        }, "required": ["path", "content"]}
    }},
    {"type": "function", "function": {
        "name": "create_xlsx",
        "description": "Creates an Excel file (.xlsx). rows_json can be: array of arrays [[\"Name\",\"Age\"],[\"Alice\",30]] OR array of dicts [{\"Name\":\"Alice\",\"Age\":30}]",
        "parameters": {"type": "object", "properties": {
            "path":      {"type": "string", "description": "Relative path to the .xlsx file"},
            "rows_json": {"type": "string", "description": "JSON array of rows (arrays or dicts)"}
        }, "required": ["path", "rows_json"]}
    }},
    {"type": "function", "function": {
        "name": "create_csv",
        "description": "Creates a CSV file. rows_json can be: array of arrays [[\"Name\",\"Age\"],[\"Alice\",30]] OR array of dicts [{\"Name\":\"Alice\",\"Age\":30}]",
        "parameters": {"type": "object", "properties": {
            "path":      {"type": "string", "description": "Relative path to the .csv file"},
            "rows_json": {"type": "string", "description": "JSON array of rows (arrays or dicts)"}
        }, "required": ["path", "rows_json"]}
    }},
    {"type": "function", "function": {
        "name": "csv_to_excel",
        "description": "Converts a CSV file to an Excel file (.xlsx)",
        "parameters": {"type": "object", "properties": {
            "csv_path":   {"type": "string", "description": "Relative path to the source CSV file"},
            "excel_path": {"type": "string", "description": "Relative path to the target Excel file (optional, default: same name with .xlsx)"}
        }, "required": ["csv_path"]}
    }},
    {"type": "function", "function": {
        "name": "analyse_data",
        "description": "Analyses a CSV or Excel file with pandas and returns statistics",
        "parameters": {"type": "object", "properties": {
            "path": {"type": "string", "description": "Relative path to the CSV or Excel file"}
        }, "required": ["path"]}
    }},
    {"type": "function", "function": {
        "name": "create_chart",
        "description": "Creates a chart from a CSV or Excel file and saves it as PNG",
        "parameters": {"type": "object", "properties": {
            "path":        {"type": "string", "description": "Relative path to the data file (CSV or Excel)"},
            "chart_type":  {"type": "string", "description": "Chart type: line, bar, scatter, pie, hist, boxplot"},
            "x_col":       {"type": "string", "description": "Name of the x-axis column"},
            "y_col":       {"type": "string", "description": "Name of the y-axis column"},
            "title":       {"type": "string", "description": "Chart title (optional)"},
            "output_path": {"type": "string", "description": "Output path for the PNG file (default: chart.png)"}
        }, "required": ["path", "chart_type", "x_col", "y_col"]}
    }},
    {"type": "function", "function": {
        "name": "create_ascii_art",
        "description": "Generates ASCII art for a description. Use this whenever the user asks to draw, design, paint, sketch, or create visual/text art.",
        "parameters": {"type": "object", "properties": {
            "description": {"type": "string", "description": "What to draw as ASCII art (e.g. 'a heart with WO AI NI inside')"}
        }, "required": ["description"]}
    }},
]

# ── Tool-call fallback (for models that output JSON text instead of tool_calls) ─
@dataclass
class _Fn:
    name: str
    arguments: dict

@dataclass
class _TC:
    function: _Fn

def _extract_json_objects(text: str) -> list:
    """Extract all JSON objects {…} from text, handling nesting and pretty-printing."""
    results = []
    i = 0
    while i < len(text):
        if text[i] == '{':
            depth = 0
            for j in range(i, len(text)):
                if text[j] == '{':
                    depth += 1
                elif text[j] == '}':
                    depth -= 1
                    if depth == 0:
                        try:
                            obj = json.loads(text[i:j + 1])
                            if isinstance(obj, dict) and "name" in obj:
                                results.append(obj)
                        except (json.JSONDecodeError, ValueError):
                            pass
                        i = j
                        break
        i += 1
    return results

def _get_tool_calls(msg) -> list:
    """Return tool calls from the message, with fallback for models that
    serialize tool calls as plain JSON in content instead of tool_calls."""
    if msg.tool_calls:
        return msg.tool_calls
    content = (msg.content or "").strip()
    if not content:
        return []
    # Try 1: entire content is valid JSON (single object or array)
    try:
        data = json.loads(content)
        if isinstance(data, dict) and "name" in data:
            data = [data]
        if isinstance(data, list):
            calls = [_TC(_Fn(d["name"], d.get("arguments", {})))
                     for d in data if "name" in d]
            if calls:
                return calls
    except (json.JSONDecodeError, TypeError, KeyError):
        pass
    # Try 2: JSON objects embedded anywhere in the text (handles pretty-printed JSON)
    objects = _extract_json_objects(content)
    if objects:
        return [_TC(_Fn(o["name"], o.get("arguments", {}))) for o in objects]
    # Try 3: newline-delimited JSON (one compact object per line)
    calls = []
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
            if isinstance(obj, dict) and "name" in obj:
                calls.append(_TC(_Fn(obj["name"], obj.get("arguments", {}))))
        except (json.JSONDecodeError, TypeError, KeyError):
            pass
    return calls

# ── Conversation history (in-memory) ─────────────────────────────────────────
history: list[dict] = []

def _compress_messages(messages: list[dict], task: str) -> list[dict]:
    """Summarise the middle portion of messages to keep context bounded.
    Keeps the system prompt and the most recent 4 messages intact."""
    if len(messages) < 7:
        return messages
    system  = messages[0]
    recent  = messages[-4:]
    middle  = messages[1:-4]
    if not middle:
        return messages
    middle_text = "\n".join(
        f"{m['role']}: {str(m.get('content', ''))[:400]}" for m in middle
    )
    summary_msgs = [
        {"role": "system", "content": "Summarise the following agent steps in 2-3 sentences. Focus on what was done and what remains."},
        {"role": "user",   "content": middle_text},
    ]
    try:
        resp    = ollama.chat(model=MODEL, messages=summary_msgs)
        summary = resp.message.content or ""
    except Exception:
        # Fallback: plain truncation — keep last portion only
        return [system] + messages[-6:]
    return [system, {"role": "user", "content": f"[Conversation summary: {summary}]"}] + recent

# ── Agent loop ────────────────────────────────────────────────────────────────
def run_agent_stream(user_message: str):
    """Generator that yields SSE-style event dicts during agent execution:
      {"type": "tool_start", "tool": name, "args": {...}}
      {"type": "tool_done",  "tool": name}
      {"type": "answer",     "text": ..., "tool_trace": [...]}
    """
    global history

    # Cap cross-turn history to prevent unbounded context growth
    if len(history) > HISTORY_MAX_MESSAGES:
        history = history[-HISTORY_MAX_MESSAGES:]

    history.append({"role": "user", "content": user_message})
    _log_message("user", user_message)

    # Fresh system prompt with relevance-scored memories for this query
    messages        = [{"role": "system", "content": _build_system(user_message)}] + history
    tool_trace      = []
    tool_call_count = 0

    for _ in range(10):
        # Dynamic tool selection — only send tools relevant to this query
        active_tools = _select_tools(user_message)
        response     = ollama.chat(model=MODEL, messages=messages, tools=active_tools,
                                   options={"num_predict": 1024})
        msg          = response.message
        tool_calls   = _get_tool_calls(msg)

        if not tool_calls:
            answer = msg.content or ""
            history.append({"role": "assistant", "content": answer})
            _log_message("assistant", answer)
            yield {"type": "answer", "text": answer, "tool_trace": tool_trace}
            return

        # Append assistant message with tool calls
        assistant_entry = {
            "role": "assistant",
            "content": msg.content or "",
            "tool_calls": [
                {"function": {"name": tc.function.name, "arguments": tc.function.arguments}}
                for tc in tool_calls
            ],
        }
        messages.append(assistant_entry)
        history.append(assistant_entry)

        # Execute each tool
        for tc in tool_calls:
            fn   = TOOL_MAP.get(tc.function.name)
            args = tc.function.arguments or {}

            yield {"type": "tool_start", "tool": tc.function.name, "args": args}

            # Web-search confirmation gate
            if tc.function.name == "web_search":
                global _search_confirm_event, _search_confirm_approved
                _search_confirm_approved = False
                _search_confirm_event.clear()
                yield {"type": "confirm_search", "query": args.get("query", "")}
                granted = _search_confirm_event.wait(timeout=SEARCH_CONFIRM_TIMEOUT)
                if not granted or not _search_confirm_approved:
                    result = "Websuche wurde vom Nutzer abgebrochen."
                    tool_trace.append({"tool": tc.function.name, "args": args, "result": result})
                    tool_entry = {"role": "tool", "content": result}
                    messages.append(tool_entry)
                    history.append(tool_entry)
                    tool_call_count += 1
                    continue

            try:
                result = fn(**args) if fn else f"Unbekanntes Tool: {tc.function.name}"
            except Exception as e:
                result = f"Fehler beim Ausführen von '{tc.function.name}': {e}"

            tool_trace.append({
                "tool":   tc.function.name,
                "args":   args,
                "result": result[:800],
            })

            yield {"type": "tool_done", "tool": tc.function.name}

            tool_entry = {"role": "tool", "content": result}
            messages.append(tool_entry)
            history.append(tool_entry)
            tool_call_count += 1

        # Context compression: summarise middle history every N tool calls
        if tool_call_count > 0 and tool_call_count % CONTEXT_COMPRESS_AFTER == 0:
            messages = _compress_messages(messages, user_message)

    yield {"type": "answer", "text": "⚠️ Maximale Iterationen erreicht.", "tool_trace": tool_trace}


def run_agent(user_message: str) -> dict:
    """Synchronous wrapper around run_agent_stream (used by /chat and /memorize)."""
    for event in run_agent_stream(user_message):
        if event["type"] == "answer":
            return {"answer": event["text"], "tool_trace": event["tool_trace"]}
    return {"answer": "⚠️ Maximale Iterationen erreicht.", "tool_trace": []}

# ── FastAPI app ───────────────────────────────────────────────────────────────
app = FastAPI()
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:8000",
        "http://127.0.0.1:8000",
        f"http://{_local_ip()}:8000",
    ],
    allow_methods=["GET", "POST"],
    allow_headers=["Content-Type"],
)

from starlette.middleware.base import BaseHTTPMiddleware
from starlette.requests import Request

class SecurityHeadersMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        response = await call_next(request)
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["Referrer-Policy"] = "no-referrer"
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; script-src 'unsafe-inline'; style-src 'unsafe-inline'; "
            "img-src 'self' data:; connect-src 'self'; frame-ancestors 'none';"
        )
        return response

app.add_middleware(SecurityHeadersMiddleware)

class MessageRequest(BaseModel):
    message: str = ""

    class Config:
        @staticmethod
        def json_schema_extra(schema: dict) -> None:
            schema["properties"]["message"]["maxLength"] = 50000

    def __init__(self, **data):
        super().__init__(**data)
        if len(self.message) > 50000:
            raise ValueError("Message too long (max 50000 characters)")

class SetupRequest(BaseModel):
    agent_name: str = ""
    user_goal:  str = ""

    def __init__(self, **data):
        super().__init__(**data)
        if len(self.agent_name) > 40:
            raise ValueError("Agent name too long (max 40 characters)")
        if len(self.user_goal) > 2000:
            raise ValueError("User goal too long (max 2000 characters)")

@app.post("/chat")
async def chat(req: MessageRequest):
    async with _request_lock:
        return JSONResponse(run_agent(req.message))

@app.post("/chat-stream")
async def chat_stream(req: MessageRequest):
    """SSE endpoint — streams tool_start / tool_done / answer events."""
    q: _queue.Queue = _queue.Queue()

    def run() -> None:
        try:
            for event in run_agent_stream(req.message):
                q.put(event)
        except Exception as e:
            q.put({"type": "error", "text": str(e)})
        finally:
            q.put(None)  # sentinel

    async def generate():
        loop = asyncio.get_event_loop()
        fut  = loop.run_in_executor(None, run)
        while True:
            try:
                event = q.get_nowait()
            except _queue.Empty:
                await asyncio.sleep(0.02)
                continue
            if event is None:
                break
            yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
        await fut

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )

@app.post("/reset")
async def reset():
    global history
    async with _request_lock:
        history = []
    return {"ok": True}

@app.get("/status")
async def status():
    return {"model": MODEL, "root": str(ROOT)}

@app.get("/config")
async def config():
    return {
        "agent_name": _cfg_get("agent_name"),
        "user_goal":  _cfg_get("user_goal"),
    }

@app.post("/setup")
async def setup(req: SetupRequest):
    _cfg_set("agent_name", req.agent_name.strip())
    _cfg_set("user_goal",  req.user_goal.strip())
    return {"ok": True}

class SetRootRequest(BaseModel):
    path: str = ""
    def __init__(self, **data):
        super().__init__(**data)
        if len(self.path) > 500:
            raise ValueError("Path too long")

@app.post("/set-root")
async def set_root(req: SetRootRequest):
    global ROOT, history
    new_path = Path(req.path.strip()).expanduser().resolve()
    if not new_path.is_dir():
        return JSONResponse({"error": "Directory not found"}, status_code=400)
    # Security: only allow subdirectories of the original root, never the whole filesystem
    if new_path != ORIGINAL_ROOT and ORIGINAL_ROOT not in new_path.parents:
        return JSONResponse({"error": "Path must be within the original root directory"}, status_code=403)
    ROOT = new_path
    history = []
    return {"ok": True, "root": str(ROOT)}

@app.get("/file")
async def serve_file(path: str):
    t = _safe(path)
    if not t or not t.is_file():
        return JSONResponse({"error": "File not found"}, status_code=404)
    return FileResponse(str(t))

MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    filename = Path(file.filename).name  # strip any path components
    if not filename:
        return JSONResponse({"error": "Invalid filename"}, status_code=400)
    dest = ROOT / filename
    # Security: only write inside root
    if not (dest == ROOT or ROOT in dest.resolve().parents):
        return JSONResponse({"error": "Invalid upload path"}, status_code=400)
    # Security: refuse to overwrite protected or existing files
    if filename in _PROTECTED_FILES:
        return JSONResponse({"error": f"'{filename}' is a protected file"}, status_code=403)
    if dest.exists():
        return JSONResponse({"error": f"'{filename}' already exists. Rename the file before uploading."}, status_code=409)
    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        return JSONResponse({"error": "File too large (max 50 MB)"}, status_code=413)
    dest.write_bytes(content)
    try:
        text_preview = content[:3000].decode("utf-8")
        is_text = True
    except UnicodeDecodeError:
        text_preview = ""
        is_text = False
    return {"path": filename, "name": file.filename, "size": len(content),
            "is_text": is_text, "preview": text_preview}

@app.get("/qrcode")
async def qrcode_endpoint():
    import qrcode as _qr
    ip  = _local_ip()
    url = f"http://{ip}:8000"
    buf = io.BytesIO()
    _qr.make(url).save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode()
    reachable = HOST == "0.0.0.0"
    return {"url": url, "qr": b64, "reachable": reachable}

@app.get("/notifications")
async def notifications():
    now = datetime.datetime.now().isoformat(sep=" ", timespec="minutes")
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            "SELECT id, title, description, due_at FROM appointments "
            "WHERE due_at <= ? AND notified = 0",
            (now,)
        ).fetchall()
        if rows:
            ids = [r[0] for r in rows]
            conn.execute(
                f"UPDATE appointments SET notified=1 WHERE id IN ({','.join('?'*len(ids))})",
                ids
            )
    return {"notifications": [{"title": r[1], "description": r[2], "due_at": r[3]} for r in rows]}

class SearchConfirmRequest(BaseModel):
    approved: bool = False

@app.post("/search-confirm")
async def search_confirm(req: SearchConfirmRequest):
    global _search_confirm_approved
    _search_confirm_approved = req.approved
    _search_confirm_event.set()
    return {"ok": True}

class MemoryUpdateRequest(BaseModel):
    content:    str   = ""
    category:   str   = "general"
    importance: float = 0.5
    def __init__(self, **data):
        super().__init__(**data)
        if len(self.content) > 2000:
            raise ValueError("Memory content too long (max 2000 chars)")
        self.importance = max(0.0, min(1.0, float(self.importance)))

@app.get("/memories")
async def get_memories_endpoint():
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            "SELECT id, created_at, category, content, importance FROM memories ORDER BY id DESC"
        ).fetchall()
    return {"memories": [
        {"id": r[0], "created_at": r[1], "category": r[2], "content": r[3], "importance": r[4] or 0.5}
        for r in rows
    ]}

@app.post("/memories")
async def add_memory_endpoint(req: MemoryUpdateRequest):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "INSERT INTO memories (content, category, importance) VALUES (?,?,?)",
            (req.content.strip(), req.category.strip(), req.importance)
        )
    return {"ok": True}

@app.put("/memories/{memory_id}")
async def update_memory_endpoint(memory_id: int, req: MemoryUpdateRequest):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "UPDATE memories SET content=?, category=?, importance=? WHERE id=?",
            (req.content.strip(), req.category.strip(), req.importance, memory_id)
        )
    return {"ok": True}

@app.delete("/memories/{memory_id}")
async def delete_memory_endpoint(memory_id: int):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM memories WHERE id=?", (memory_id,))
    return {"ok": True}

@app.post("/memorize")
async def memorize():
    """Summarise conversation history and force-save key facts to memory."""
    if not history:
        return {"answer": "No conversation to summarise yet.", "tool_trace": []}
    convo = "\n".join(
        f"{m['role']}: {str(m.get('content', ''))[:500]}"
        for m in history[-30:]
        if m.get("content")
    )
    summary_msgs = [
        {"role": "system", "content": "You are a concise summariser. Output plain text only."},
        {"role": "user",   "content":
            "Summarise the following conversation in 3–5 bullet points. "
            "Focus on facts, preferences, decisions, and context worth remembering.\n\n" + convo},
    ]
    try:
        resp    = ollama.chat(model=MODEL, messages=summary_msgs)
        summary = (resp.message.content or "").strip()
    except Exception as e:
        return {"answer": f"Error during summarisation: {e}", "tool_trace": []}
    result = save_memory(summary, category="result", importance=0.8)
    answer_text = f"✅ Conversation summarised and saved to memory:\n\n{summary}"
    _log_message("assistant", answer_text)
    return {
        "answer": answer_text,
        "tool_trace": [{"tool": "save_memory",
                        "args": {"category": "result", "importance": 0.8},
                        "result": result}],
    }

@app.get("/appointments")
async def get_appointments_endpoint():
    now = datetime.datetime.now().isoformat(sep=" ", timespec="minutes")
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            "SELECT id, due_at, title, description, notified FROM appointments ORDER BY due_at"
        ).fetchall()
    return {"appointments": [
        {"id": r[0], "due_at": r[1], "title": r[2], "description": r[3],
         "past": r[1] < now}
        for r in rows
    ]}

@app.delete("/appointments/{appt_id}")
async def delete_appointment_endpoint(appt_id: int):
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM appointments WHERE id=?", (appt_id,))
    return {"ok": True}

@app.get("/chat-log")
async def get_chat_log(limit: int = 100):
    limit = min(max(1, limit), 500)
    with sqlite3.connect(DB_PATH) as conn:
        rows = conn.execute(
            "SELECT id, created_at, role, content FROM chat_log ORDER BY id DESC LIMIT ?",
            (limit,)
        ).fetchall()
    return {"entries": [
        {"id": r[0], "created_at": r[1], "role": r[2], "content": r[3]}
        for r in rows
    ]}

@app.delete("/chat-log")
async def clear_chat_log():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM chat_log")
    return {"ok": True}

@app.get("/", response_class=HTMLResponse)
async def index():
    return HTML_PAGE

# ── Embedded Web UI ───────────────────────────────────────────────────────────
HTML_PAGE = """<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Reiseki</title>
<!-- Fonts: using system font stack to avoid external requests (data locality) -->
<!-- sepia warm-dark theme replaces previous cold-dark theme -->
<style>
  :root {
    --bg:        #ffffff;
    --surface:   #f7f5f2;
    --border:    #e2ddd5;
    --border2:   #cfc8bb;
    --text:      #1a1410;
    --muted:     #8a7a66;
    --accent:    #b87333;
    --accent2:   #9a5f1e;
    --user-bg:   #fdf8f0;
    --tool-bg:   #f5f1e8;
    --tool-text: #b87333;
    --err:       #cc3333;
    --radius:    8px;
    --mono:      'SF Mono', 'Cascadia Code', 'Fira Code', 'Consolas', 'Liberation Mono', monospace;
    --sans:      -apple-system, BlinkMacSystemFont, 'Segoe UI', system-ui, Roboto, 'Helvetica Neue', sans-serif;
  }

  * { box-sizing: border-box; margin: 0; padding: 0; }

  body {
    background: var(--bg);
    color: var(--text);
    font-family: var(--sans);
    height: 100dvh;
    display: flex;
    flex-direction: column;
    overflow: hidden;
  }

  /* ── Header ── */
  header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 12px 20px;
    border-bottom: 1px solid var(--border);
    background: var(--surface);
    flex-shrink: 0;
    gap: 12px;
  }
  .header-left {
    display: flex;
    align-items: center;
    gap: 12px;
  }
  .logo {
    font-family: var(--mono);
    font-size: 13px;
    font-weight: 600;
    color: var(--accent);
    letter-spacing: 0.05em;
    text-transform: uppercase;
  }
  .divider-v {
    width: 1px; height: 20px;
    background: var(--border2);
  }
  .meta {
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    display: flex;
    gap: 16px;
  }
  .meta span { display: flex; align-items: center; gap: 5px; }
  .meta .dot {
    width: 6px; height: 6px; border-radius: 50%;
    background: var(--accent);
    box-shadow: 0 0 6px var(--accent);
    animation: pulse 2s infinite;
  }
  @keyframes pulse {
    0%, 100% { opacity: 1; }
    50%       { opacity: 0.4; }
  }
  .btn-reset {
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    background: none;
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    padding: 5px 12px;
    cursor: pointer;
    transition: all 0.15s;
  }
  .btn-reset:hover { color: var(--err); border-color: var(--err); }

  /* ── Chat area ── */
  #chat {
    flex: 1;
    overflow-y: auto;
    padding: 24px 20px;
    display: flex;
    flex-direction: column;
    gap: 20px;
    scroll-behavior: smooth;
  }
  #chat::-webkit-scrollbar { width: 4px; }
  #chat::-webkit-scrollbar-track { background: transparent; }
  #chat::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }

  .msg { display: flex; flex-direction: column; gap: 6px; max-width: 820px; width: 100%; }
  .msg.user  { align-self: flex-end; align-items: flex-end; }
  .msg.agent { align-self: flex-start; align-items: flex-start; }

  .msg-label {
    font-family: var(--mono);
    font-size: 10px;
    color: var(--muted);
    letter-spacing: 0.08em;
    text-transform: uppercase;
    padding: 0 4px;
  }

  .bubble {
    padding: 12px 16px;
    border-radius: var(--radius);
    line-height: 1.6;
    font-size: 14px;
    white-space: pre-wrap;
    word-break: break-word;
  }
  .msg.user  .bubble { background: var(--user-bg); border: 1px solid var(--border2); color: var(--text); }
  .msg.agent .bubble { background: var(--surface); border: 1px solid var(--border); color: var(--text); }

  /* ── Tool trace ── */
  .tool-trace { display: flex; flex-direction: column; gap: 4px; width: 100%; }

  .tool-item {
    border: 1px solid var(--border);
    border-radius: var(--radius);
    overflow: hidden;
    background: var(--tool-bg);
  }
  .tool-header {
    display: flex;
    align-items: center;
    gap: 8px;
    padding: 7px 12px;
    cursor: pointer;
    user-select: none;
    transition: background 0.1s;
  }
  .tool-header:hover { background: rgba(184,115,51,0.06); }
  .tool-icon { font-size: 12px; }
  .tool-name {
    font-family: var(--mono);
    font-size: 11px;
    color: var(--tool-text);
    font-weight: 600;
  }
  .tool-args {
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    flex: 1;
    overflow: hidden;
    text-overflow: ellipsis;
    white-space: nowrap;
  }
  .tool-chevron {
    font-size: 9px;
    color: var(--muted);
    transition: transform 0.2s;
    margin-left: auto;
  }
  .tool-item.open .tool-chevron { transform: rotate(90deg); }

  .tool-result {
    display: none;
    padding: 10px 14px;
    border-top: 1px solid var(--border);
    font-family: var(--mono);
    font-size: 11px;
    color: #5a6070;
    white-space: pre-wrap;
    word-break: break-all;
    line-height: 1.7;
    max-height: 300px;
    overflow-y: auto;
  }
  .tool-item.open .tool-result { display: block; }

  /* ── Loader ── */
  .loader {
    display: flex;
    align-items: center;
    gap: 10px;
    padding: 12px 16px;
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    font-family: var(--mono);
    font-size: 12px;
    color: var(--muted);
  }
  .loader-dots { display: flex; gap: 4px; }
  .loader-dots span {
    width: 4px; height: 4px; border-radius: 50%;
    background: var(--accent);
    animation: dot-bounce 1.2s infinite;
  }
  .loader-dots span:nth-child(2) { animation-delay: 0.2s; }
  .loader-dots span:nth-child(3) { animation-delay: 0.4s; }
  @keyframes dot-bounce {
    0%, 60%, 100% { transform: translateY(0); opacity: 0.4; }
    30%            { transform: translateY(-5px); opacity: 1; }
  }

  /* ── Empty state ── */
  .empty {
    flex: 1;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    gap: 12px;
    color: var(--muted);
    text-align: center;
  }
  .empty-glyph { font-size: 36px; opacity: 0.3; }
  .empty-title { font-family: var(--mono); font-size: 13px; letter-spacing: 0.06em; }
  .empty-sub   { font-size: 12px; line-height: 1.6; max-width: 340px; }
  .hint-pills  { display: flex; flex-wrap: wrap; gap: 8px; justify-content: center; margin-top: 8px; }
  .pill {
    font-family: var(--mono);
    font-size: 11px;
    padding: 6px 12px;
    border: 1px solid var(--border2);
    border-radius: 20px;
    cursor: pointer;
    transition: all 0.15s;
    color: var(--muted);
  }
  .pill:hover { border-color: var(--accent); color: var(--accent); background: rgba(184,115,51,0.08); }
  .btn-guide {
    display: flex;
    flex-wrap: wrap;
    gap: 6px;
    justify-content: center;
    margin-top: 4px;
    font-size: 11px;
  }
  .btn-guide span {
    background: rgba(184,115,51,0.06);
    border: 1px solid var(--border);
    border-radius: 10px;
    padding: 3px 10px;
    color: var(--muted);
    font-family: var(--mono);
  }
  .lang-toggle {
    font-family: var(--mono);
    font-size: 10px;
    letter-spacing: 0.05em;
    padding: 4px 8px;
  }

  /* ── Input bar ── */
  footer {
    padding: 14px 20px;
    border-top: 1px solid var(--border);
    background: var(--surface);
    flex-shrink: 0;
  }
  .input-row {
    display: flex;
    gap: 10px;
    align-items: flex-end;
    max-width: 860px;
    margin: 0 auto;
  }
  textarea {
    flex: 1;
    background: var(--bg);
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    color: var(--text);
    font-family: var(--sans);
    font-size: 14px;
    padding: 10px 14px;
    resize: none;
    min-height: 44px;
    max-height: 160px;
    outline: none;
    transition: border-color 0.15s;
    line-height: 1.5;
  }
  textarea::placeholder { color: var(--muted); }
  textarea:focus { border-color: var(--accent2); }

  .btn-send {
    width: 44px; height: 44px;
    background: var(--accent);
    border: none;
    border-radius: var(--radius);
    cursor: pointer;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
    transition: all 0.15s;
    color: #ffffff;
    font-size: 18px;
  }
  .btn-send:hover:not(:disabled) { background: var(--accent2); transform: scale(1.05); }
  .btn-send:disabled { opacity: 0.35; cursor: not-allowed; transform: none; }

  /* ── Modals (shared) ── */
  .modal-overlay {
    display: none;
    position: fixed; inset: 0;
    background: rgba(0,0,0,0.7);
    z-index: 100;
    align-items: center;
    justify-content: center;
  }
  .modal-overlay.open { display: flex; }
  .modal {
    background: var(--surface);
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    padding: 28px 32px;
    width: min(480px, 92vw);
    display: flex;
    flex-direction: column;
    gap: 16px;
  }
  .modal h2 {
    font-family: var(--mono);
    font-size: 13px;
    letter-spacing: 0.06em;
    color: var(--accent);
    text-transform: uppercase;
    margin: 0;
  }
  .modal p { font-size: 13px; color: var(--muted); line-height: 1.6; margin: 0; }
  .modal label {
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    letter-spacing: 0.05em;
    display: flex;
    flex-direction: column;
    gap: 6px;
  }
  .modal input, .modal textarea {
    background: var(--bg);
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    color: var(--text);
    font-family: var(--sans);
    font-size: 13px;
    padding: 8px 12px;
    outline: none;
    transition: border-color 0.15s;
    resize: none;
  }
  .modal input:focus, .modal textarea:focus { border-color: var(--accent); }
  .modal-actions { display: flex; justify-content: flex-end; gap: 10px; }
  .btn-primary {
    font-family: var(--mono); font-size: 12px;
    background: var(--accent); color: var(--bg);
    border: none; border-radius: var(--radius);
    padding: 8px 20px; cursor: pointer; transition: background 0.15s;
  }
  .btn-primary:hover { background: var(--accent2); }
  .btn-secondary {
    font-family: var(--mono); font-size: 12px;
    background: none; color: var(--muted);
    border: 1px solid var(--border2); border-radius: var(--radius);
    padding: 8px 16px; cursor: pointer; transition: all 0.15s;
  }
  .btn-secondary:hover { color: var(--text); border-color: var(--text); }

  /* ── QR modal specifics ── */
  .qr-img { display: block; margin: 0 auto; border-radius: 6px; max-width: 220px; }
  .qr-url {
    font-family: var(--mono); font-size: 12px; color: var(--accent);
    text-align: center; word-break: break-all;
  }

  /* ── Toast notifications ── */
  #toast-container {
    position: fixed; top: 16px; right: 16px;
    display: flex; flex-direction: column; gap: 8px;
    z-index: 200; pointer-events: none;
  }
  .toast {
    background: var(--surface);
    border: 1px solid var(--accent);
    border-radius: var(--radius);
    padding: 12px 16px;
    max-width: 300px;
    font-size: 13px;
    line-height: 1.5;
    pointer-events: all;
    animation: slide-in 0.25s ease;
  }
  .toast-title { font-family: var(--mono); font-size: 11px; color: var(--accent); margin-bottom: 4px; }
  @keyframes slide-in {
    from { opacity: 0; transform: translateX(20px); }
    to   { opacity: 1; transform: translateX(0); }
  }

  /* ── Code canvas ── */
  .code-block {
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    overflow: hidden;
    margin: 6px 0;
    background: var(--tool-bg);
  }
  .code-header {
    display: flex;
    align-items: center;
    justify-content: space-between;
    padding: 5px 10px;
    background: var(--border);
    font-family: var(--mono);
    font-size: 10px;
  }
  .code-lang { color: var(--accent); font-weight: 600; text-transform: uppercase; letter-spacing: 0.05em; }
  .copy-btn {
    background: none;
    border: 1px solid var(--border2);
    border-radius: 4px;
    color: var(--muted);
    font-family: var(--mono);
    font-size: 10px;
    padding: 2px 8px;
    cursor: pointer;
    transition: all 0.15s;
  }
  .copy-btn:hover { color: var(--accent); border-color: var(--accent); }
  .code-pre {
    margin: 0;
    padding: 12px 14px;
    font-family: var(--mono);
    font-size: 12px;
    line-height: 1.6;
    overflow-x: auto;
    white-space: pre;
    color: var(--text);
  }
  .bubble .text-part { white-space: pre-wrap; word-break: break-word; }

  /* ── Chart / inline image ── */
  .chat-img {
    max-width: 100%;
    max-height: 420px;
    border-radius: var(--radius);
    border: 1px solid var(--border2);
    margin-top: 8px;
    display: block;
    cursor: zoom-in;
  }
  .chat-img.fullsize { max-height: none; cursor: zoom-out; }

  /* ── File upload drag-over ── */
  #chat.drag-over {
    outline: 2px dashed var(--accent);
    outline-offset: -6px;
    background: rgba(184,115,51,0.04);
  }

  /* ── File / memorize buttons ── */
  .btn-file {
    width: 44px; height: 44px;
    background: none;
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    cursor: pointer;
    display: flex; align-items: center; justify-content: center;
    flex-shrink: 0;
    font-size: 18px;
    color: var(--muted);
    transition: all 0.15s;
  }
  .btn-file:hover { border-color: var(--accent); color: var(--accent); }
  .btn-memorize {
    height: 44px;
    padding: 0 12px;
    background: none;
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    cursor: pointer;
    display: flex; align-items: center; gap: 5px;
    flex-shrink: 0;
    font-family: var(--mono);
    font-size: 11px;
    color: var(--muted);
    transition: all 0.15s;
    white-space: nowrap;
  }
  .btn-memorize:hover:not(:disabled) { border-color: var(--accent); color: var(--accent); }
  .btn-memorize:disabled { opacity: 0.35; cursor: not-allowed; }

  /* ── Appointment cards ── */
  .appt-card {
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 10px 12px;
    display: flex;
    align-items: flex-start;
    gap: 10px;
    background: var(--bg);
  }
  .appt-card.past { opacity: 0.55; }
  .appt-icon { font-size: 16px; flex-shrink: 0; padding-top: 1px; }
  .appt-body { flex: 1; display: flex; flex-direction: column; gap: 2px; }
  .appt-title { font-size: 13px; font-weight: 600; color: var(--text); }
  .appt-date  { font-family: var(--mono); font-size: 10px; color: var(--accent); }
  .appt-desc  { font-size: 12px; color: var(--muted); line-height: 1.5; }

  /* ── Memory modal ── */
  .modal.memory-modal {
    width: min(700px, 95vw);
    max-height: 82vh;
    overflow: hidden;
    display: flex;
    flex-direction: column;
  }
  .memory-list {
    overflow-y: auto;
    flex: 1;
    display: flex;
    flex-direction: column;
    gap: 8px;
    padding-right: 4px;
    min-height: 60px;
  }
  .memory-list::-webkit-scrollbar { width: 4px; }
  .memory-list::-webkit-scrollbar-thumb { background: var(--border2); border-radius: 2px; }
  .memory-card {
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 10px 12px;
    display: flex;
    flex-direction: column;
    gap: 6px;
    background: var(--bg);
  }
  .memory-card-row {
    display: flex;
    gap: 8px;
    align-items: center;
  }
  .category-badge {
    font-family: var(--mono);
    font-size: 10px;
    padding: 2px 8px;
    border-radius: 10px;
    background: rgba(184,115,51,0.12);
    color: var(--accent);
    font-weight: 600;
    flex-shrink: 0;
  }
  .importance-bar {
    height: 4px;
    border-radius: 2px;
    background: var(--border2);
    flex: 1;
    overflow: hidden;
  }
  .importance-fill { height: 100%; background: var(--accent); border-radius: 2px; }
  .mem-importance {
    width: 52px;
    font-size: 11px;
    padding: 2px 4px;
    border: 1px solid var(--border2);
    border-radius: 4px;
    background: var(--bg);
    color: var(--text);
    text-align: center;
  }
  .mem-category {
    font-size: 11px;
    padding: 3px 6px;
    border: 1px solid var(--border2);
    border-radius: 4px;
    background: var(--bg);
    color: var(--muted);
    width: 130px;
    font-family: var(--mono);
  }
  .mem-content {
    font-size: 13px;
    line-height: 1.5;
    color: var(--text);
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 4px;
    padding: 6px 8px;
    width: 100%;
    resize: vertical;
    min-height: 40px;
    font-family: var(--sans);
    outline: none;
    transition: border-color 0.15s;
  }
  .mem-content:focus { border-color: var(--accent); }
  .memory-meta { font-family: var(--mono); font-size: 10px; color: var(--muted); }
  .btn-delete {
    background: none;
    border: 1px solid var(--border2);
    border-radius: 4px;
    color: var(--muted);
    font-size: 12px;
    padding: 3px 8px;
    cursor: pointer;
    transition: all 0.15s;
    flex-shrink: 0;
  }
  .btn-delete:hover { color: var(--err); border-color: var(--err); }
  .btn-save-mem {
    background: none;
    border: 1px solid var(--border2);
    border-radius: 4px;
    color: var(--muted);
    font-size: 11px;
    font-family: var(--mono);
    padding: 3px 10px;
    cursor: pointer;
    transition: all 0.15s;
  }
  .btn-save-mem:hover { color: var(--accent); border-color: var(--accent); }

  /* ── Search confirmation modal ── */
  .search-confirm-query {
    font-family: var(--mono);
    font-size: 13px;
    background: var(--tool-bg);
    border: 1px solid var(--border2);
    border-radius: var(--radius);
    padding: 10px 14px;
    color: var(--text);
    word-break: break-word;
    line-height: 1.5;
  }
  .search-confirm-query::before {
    content: '🔍  ';
  }

  /* ── History modal entries ── */
  .history-entry {
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 10px 12px;
    display: flex;
    flex-direction: column;
    gap: 4px;
    background: var(--bg);
  }
  .history-entry.user-entry  { border-left: 3px solid var(--border2); }
  .history-entry.agent-entry { border-left: 3px solid var(--accent); }
  .history-role {
    font-family: var(--mono);
    font-size: 10px;
    color: var(--muted);
    display: flex;
    justify-content: space-between;
    align-items: center;
  }
  .history-role .role-label { color: var(--accent); font-weight: 600; }
  .history-content {
    font-size: 13px;
    color: var(--text);
    white-space: pre-wrap;
    word-break: break-word;
    line-height: 1.5;
    max-height: 160px;
    overflow-y: auto;
  }

  /* ── Root-path clickable ── */
  #root-path {
    cursor: pointer;
    border-bottom: 1px dashed var(--border2);
    transition: color 0.15s;
  }
  #root-path:hover { color: var(--accent); }

  /* ── QR reachability warning ── */
  .qr-warning {
    background: #fff3cd;
    border: 1px solid #ffc107;
    border-radius: var(--radius);
    padding: 8px 12px;
    font-size: 12px;
    color: #856404;
    line-height: 1.5;
  }
</style>
</head>
<body>

<header>
  <div class="header-left">
    <span class="logo">⬡ <span id="agent-label">Reiseki</span></span>
    <div class="divider-v"></div>
    <div class="meta" id="meta">
      <span><span class="dot"></span><span id="model-name">–</span></span>
      <span>📁 <span id="root-path" onclick="openDirModal()" title="Change directory">–</span></span>
    </div>
  </div>
  <div style="display:flex;gap:8px;align-items:center;">
    <button class="btn-reset lang-toggle" id="lang-btn" onclick="toggleLang()">DE</button>
    <button class="btn-reset" onclick="openQr()" data-i18n-title="title_qr" title="QR code for smartphone">📱</button>
    <button class="btn-reset" onclick="openMemory()" data-i18n-title="title_memory" title="Saved knowledge">🧠</button>
    <button class="btn-reset" onclick="openAppointments()" data-i18n-title="title_appointments" title="Appointments">📅</button>
    <button class="btn-reset" onclick="openHistory()" data-i18n-title="title_history" title="Conversation history">📜</button>
    <button class="btn-reset" onclick="openSetup()" title="⚙ Setup">⚙</button>
    <button class="btn-reset" onclick="resetChat()">↺ Reset</button>
  </div>
</header>

<!-- ── Appointments Modal ── -->
<div class="modal-overlay" id="appt-modal">
  <div class="modal memory-modal">
    <h2 data-i18n="appt_title">📅 Appointments</h2>
    <p data-i18n="appt_desc">All appointments saved by the agent. Past appointments are greyed out.</p>
    <div class="memory-list" id="appt-list"></div>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="closeModal('appt-modal')" data-i18n="btn_close">Close</button>
    </div>
  </div>
</div>

<!-- ── Memory Modal ── -->
<div class="modal-overlay" id="memory-modal">
  <div class="modal memory-modal">
    <div style="display:flex;justify-content:space-between;align-items:center">
      <h2 data-i18n="mem_title">🧠 Saved Knowledge</h2>
      <button class="btn-primary" style="font-size:11px;padding:5px 14px" onclick="addNewMemory()" data-i18n="btn_new">+ New</button>
    </div>
    <p data-i18n="mem_desc">All agent memories. Content, category and importance (0–1) are directly editable.</p>
    <div class="memory-list" id="memory-list"></div>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="closeModal('memory-modal')" data-i18n="btn_close">Close</button>
    </div>
  </div>
</div>

<!-- ── Search Confirmation Modal ── -->
<div class="modal-overlay" id="search-confirm-modal">
  <div class="modal">
    <h2 data-i18n="search_title">🔍 Confirm web search</h2>
    <p data-i18n="search_desc">The agent wants to send the following query to DuckDuckGo:</p>
    <div class="search-confirm-query" id="search-confirm-query"></div>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="respondSearch(false)" data-i18n="btn_cancel">Cancel</button>
      <button class="btn-primary" onclick="respondSearch(true)" data-i18n="btn_search">Search</button>
    </div>
  </div>
</div>

<!-- ── History Modal ── -->
<div class="modal-overlay" id="history-modal">
  <div class="modal memory-modal">
    <div style="display:flex;justify-content:space-between;align-items:center">
      <h2 data-i18n="history_title">📜 Conversation History</h2>
      <button class="btn-secondary" style="font-size:11px;padding:5px 14px;color:var(--err);border-color:var(--err)" onclick="clearHistory()" data-i18n="btn_delete_history">Delete</button>
    </div>
    <p data-i18n="history_desc">All saved messages (newest first). Persisted across sessions.</p>
    <div class="memory-list" id="history-list"></div>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="closeModal('history-modal')" data-i18n="btn_close">Close</button>
    </div>
  </div>
</div>

<!-- ── Directory Modal ── -->
<div class="modal-overlay" id="dir-modal">
  <div class="modal">
    <h2 data-i18n="dir_title">📁 Change directory</h2>
    <p data-i18n="dir_desc">Enter the absolute path to the new working directory. Chat history will be reset.</p>
    <label><span data-i18n="label_dir">DIRECTORY</span>
      <input type="text" id="dir-input" data-i18n-placeholder="dir_placeholder" placeholder="/home/user/my-project" style="width:100%">
    </label>
    <div id="dir-error" style="color:var(--err);font-size:12px;display:none"></div>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="closeModal('dir-modal')" data-i18n="btn_cancel">Cancel</button>
      <button class="btn-primary" onclick="saveDir()" data-i18n="btn_switch">Switch</button>
    </div>
  </div>
</div>

<!-- ── QR Code Modal ── -->
<div class="modal-overlay" id="qr-modal">
  <div class="modal">
    <h2 data-i18n="qr_title">📱 Open on home network</h2>
    <p data-i18n="qr_desc">Scan the QR code with your smartphone — both devices must be on the same Wi-Fi.</p>
    <div id="qr-warning" class="qr-warning" style="display:none" data-i18n-html="qr_warning">
      ⚠️ The server is only running on <b>127.0.0.1</b> — smartphone cannot reach it.<br>
      Start with <code>AGENT_HOST=0.0.0.0 python agent.py</code> for network access.
    </div>
    <img class="qr-img" id="qr-img" src="" alt="QR Code">
    <div class="qr-url" id="qr-url"></div>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="closeModal('qr-modal')" data-i18n="btn_close">Close</button>
    </div>
  </div>
</div>

<!-- ── Setup Modal ── -->
<div class="modal-overlay" id="setup-modal">
  <div class="modal">
    <h2 data-i18n="setup_title">⚙ Configure agent</h2>
    <p data-i18n="setup_desc">Give your agent a name and describe your goal. This info is saved permanently.</p>
    <label><span data-i18n="label_agent_name">AGENT NAME</span>
      <input type="text" id="setup-name" data-i18n-placeholder="name_placeholder" placeholder="e.g. ARIA" maxlength="40">
    </label>
    <label><span data-i18n="label_goal">YOUR GOAL / CONTEXT</span>
      <textarea id="setup-goal" rows="3" data-i18n-placeholder="goal_placeholder" placeholder="e.g. I am a developer working on Python projects."></textarea>
    </label>
    <div class="modal-actions">
      <button class="btn-secondary" onclick="closeModal('setup-modal')" data-i18n="btn_cancel">Cancel</button>
      <button class="btn-primary" onclick="saveSetup()" data-i18n="btn_save">Save</button>
    </div>
  </div>
</div>

<!-- ── Toast container ── -->
<div id="toast-container"></div>

<div id="chat">
  <div class="empty" id="empty">
    <div class="empty-glyph">◈</div>
    <div class="empty-title" data-i18n="empty_title">LOCAL ASSISTENT</div>
    <div class="empty-sub">
      <span data-i18n="empty_sub">Your personal assistant for files, web search, appointments and more.</span>
      <div class="btn-guide">
        <span data-i18n="btn_guide_qr">📱 Smartphone access</span>
        <span data-i18n="btn_guide_memory">🧠 View &amp; edit memories</span>
        <span data-i18n="btn_guide_appt">📅 Manage appointments</span>
        <span data-i18n="btn_guide_history">📜 Browse conversation log</span>
      </div>
    </div>
    <div class="hint-pills">
      <div class="pill" data-i18n="pill_list" data-hint="hint_list" onclick="sendHint(STRINGS[currentLang].hint_list)">List files</div>
      <div class="pill" data-i18n="pill_create" data-hint="hint_create" onclick="sendHint(STRINGS[currentLang].hint_create)">Create file</div>
      <div class="pill" data-i18n="pill_read" data-hint="hint_read" onclick="sendHint(STRINGS[currentLang].hint_read)">Read file</div>
      <div class="pill" data-i18n="pill_search" data-hint="hint_search" onclick="sendHint(STRINGS[currentLang].hint_search)">Web search</div>
    </div>
  </div>
</div>

<footer>
  <div class="input-row">
    <input type="file" id="file-input" style="display:none" onchange="handleFileSelect(event)">
    <button class="btn-file" onclick="document.getElementById('file-input').click()" data-i18n-title="title_attach" title="Attach file">+</button>
    <textarea id="input" data-i18n-placeholder="input_placeholder" placeholder="Type a message… (Enter = Send, Shift+Enter = new line)"
              rows="1" onkeydown="handleKey(event)" oninput="autoResize(this)"></textarea>
    <button class="btn-memorize" id="memorize-btn" onclick="triggerMemorize()" data-i18n-title="title_memorize" title="Summarize conversation and save to memory">🧠 <span data-i18n="btn_memorize">Memorize</span></button>
    <button class="btn-send" id="send-btn" onclick="sendMessage()">↑</button>
  </div>
</footer>

<script>
  const chat    = document.getElementById('chat');
  const input   = document.getElementById('input');
  const sendBtn = document.getElementById('send-btn');
  const empty   = document.getElementById('empty');
  let   busy    = false;
  let   agentName = 'Agent';

  // ── i18n ──────────────────────────────────────────────────────────────────
  let currentLang = localStorage.getItem('lang') || 'en';

  const STRINGS = {
    en: {
      // Header tooltips
      title_qr:           'QR code for smartphone',
      title_memory:       'Saved knowledge',
      title_appointments: 'Appointments',
      title_history:      'Conversation history',
      // Empty state
      empty_title:        'LOCAL ASSISTENT',
      empty_sub:          'Your personal assistant for files, web search, appointments and more.',
      btn_guide_qr:       '📱 Smartphone access',
      btn_guide_memory:   '🧠 View & edit memories',
      btn_guide_appt:     '📅 Manage appointments',
      btn_guide_history:  '📜 Browse conversation log',
      // Pills
      pill_list:   'List files',
      pill_create: 'Create file',
      pill_read:   'Read file',
      pill_search: 'Web search',
      hint_list:   'Show me what is in the root directory',
      hint_create: 'Create a test.txt with Hello World',
      hint_read:   'What does the README.md say?',
      hint_search: 'Search the web for recent news',
      // Footer
      input_placeholder: 'Type a message\\u2026 (Enter = Send, Shift+Enter = new line)',
      title_memorize:    'Summarize conversation and save to memory',
      title_attach:      'Attach file',
      btn_memorize:      'Memorize',
      // Loader
      loader_thinking: 'Thinking\\u2026',
      // Appointments modal
      appt_title: '📅 Appointments',
      appt_desc:  'All appointments saved by the agent. Past appointments are greyed out.',
      appt_empty: 'No appointments saved.',
      btn_close:  'Close',
      // Memory modal
      mem_title:      '🧠 Saved Knowledge',
      mem_desc:       'All agent memories. Content, category and importance (0\\u20131) are directly editable.',
      mem_empty:      'No memories saved.',
      btn_new:        '+ New',
      label_category: 'Category:',
      btn_save_mem:   '\\u2713 Save',
      // Search confirm modal
      search_title: '🔍 Confirm web search',
      search_desc:  'The agent wants to send the following query to DuckDuckGo:',
      btn_cancel:   'Cancel',
      btn_search:   'Search',
      // History modal
      history_title:          '📜 Conversation History',
      history_desc:           'All saved messages (newest first). Persisted across sessions.',
      history_empty:          'No conversation history available.',
      btn_delete_history:     'Delete',
      confirm_clear_history:  'Delete entire conversation history?',
      // Dir modal
      dir_title:       '📁 Change directory',
      dir_desc:        'Enter the absolute path to the new working directory. Chat history will be reset.',
      label_dir:       'DIRECTORY',
      dir_placeholder: '/home/user/my-project',
      btn_switch:      'Switch',
      // QR modal
      qr_title: '📱 Open on home network',
      qr_desc:  'Scan the QR code with your smartphone \\u2014 both devices must be on the same Wi-Fi.',
      qr_warning: '\\u26a0\\ufe0f The server is only running on <b>127.0.0.1</b> \\u2014 smartphone cannot reach it.<br>Start with <code>AGENT_HOST=0.0.0.0 python agent.py</code> for network access.',
      // Setup modal
      setup_title:    '\\u2699 Configure agent',
      setup_desc:     'Give your agent a name and describe your goal. This info is saved permanently.',
      label_agent_name: 'AGENT NAME',
      name_placeholder: 'e.g. ARIA',
      label_goal:       'YOUR GOAL / CONTEXT',
      goal_placeholder: 'e.g. I am a developer working on Python projects.',
      btn_save: 'Save',
      // Dynamic JS strings
      role_user:            'You',
      upload_img_prefix:    '[Image uploaded: ',
      upload_file_prefix:   'File uploaded: ',
      upload_analyze:       ' \\u2014 please analyze it.',
      analyze_img:          'Analyze this image: ',
      upload_error:         '\\u26a0\\ufe0f Upload failed: ',
      reminder_prefix:      '📅 Reminder: ',
      connection_error:     '\\u26a0\\ufe0f Connection error: ',
      save_error:           '\\u26a0\\ufe0f Error saving: ',
      saved_feedback:       '\\u2713 Saved',
      new_memory_prompt:    'Enter new memory:',
      file_truncated:       '\\n[... truncated]',
      file_prefix:          'File: ',
    },
    de: {
      title_qr:           'QR-Code f\\u00fcr Smartphone',
      title_memory:       'Gespeichertes Wissen',
      title_appointments: 'Termine',
      title_history:      'Gespr\\u00e4chsverlauf',
      empty_title:        'LOCAL ASSISTENT',
      empty_sub:          'Dein pers\\u00f6nlicher Assistent f\\u00fcr Dateien, Websuche, Termine und mehr.',
      btn_guide_qr:       '📱 Smartphone-Zugriff',
      btn_guide_memory:   '🧠 Erinnerungen anzeigen & bearbeiten',
      btn_guide_appt:     '📅 Termine verwalten',
      btn_guide_history:  '📜 Gespr\\u00e4chsverlauf durchsuchen',
      pill_list:   'Dateien auflisten',
      pill_create: 'Datei erstellen',
      pill_read:   'Datei lesen',
      pill_search: 'Websuche',
      hint_list:   'Zeig mir, was im Root-Verzeichnis liegt',
      hint_create: 'Erstelle eine test.txt mit Hallo Welt',
      hint_read:   'Was steht in der README.md?',
      hint_search: 'Suche im Web nach aktuellen Neuigkeiten',
      input_placeholder: 'Nachricht eingeben\\u2026 (Enter = Senden, Shift+Enter = Zeilenumbruch)',
      title_memorize:    'Konversation zusammenfassen und in Memory speichern',
      title_attach:      'Datei anh\\u00e4ngen',
      btn_memorize:      'Merken',
      loader_thinking: 'Denkt nach\\u2026',
      appt_title: '📅 Termine',
      appt_desc:  'Alle vom Agenten gespeicherten Termine. Vergangene Termine werden ausgegraut dargestellt.',
      appt_empty: 'Keine Termine gespeichert.',
      btn_close:  'Schlie\\u00dfen',
      mem_title:      '🧠 Gespeichertes Wissen',
      mem_desc:       'Alle Erinnerungen des Agenten. Inhalte, Kategorie und Wichtigkeit (0\\u20131) sind direkt bearbeitbar.',
      mem_empty:      'Keine Erinnerungen gespeichert.',
      btn_new:        '+ Neu',
      label_category: 'Kategorie:',
      btn_save_mem:   '\\u2713 Speichern',
      search_title: '🔍 Websuche best\\u00e4tigen',
      search_desc:  'Der Agent m\\u00f6chte folgende Suchanfrage an DuckDuckGo senden:',
      btn_cancel:   'Abbrechen',
      btn_search:   'Suchen',
      history_title:          '📜 Gespr\\u00e4chsverlauf',
      history_desc:           'Alle gespeicherten Nachrichten (neueste zuerst). Wird \\u00fcber Sitzungen hinweg persistent gespeichert.',
      history_empty:          'Kein Gespr\\u00e4chsverlauf vorhanden.',
      btn_delete_history:     'L\\u00f6schen',
      confirm_clear_history:  'Gesamten Gespr\\u00e4chsverlauf l\\u00f6schen?',
      dir_title:       '📁 Verzeichnis wechseln',
      dir_desc:        'Gib den absoluten Pfad zum neuen Arbeitsverzeichnis ein. Die Chat-History wird zur\\u00fcckgesetzt.',
      label_dir:       'VERZEICHNIS',
      dir_placeholder: '/home/user/mein-projekt',
      btn_switch:      'Wechseln',
      qr_title: '📱 Im Heimnetz \\u00f6ffnen',
      qr_desc:  'Scanne den QR-Code mit deinem Smartphone \\u2014 beide Ger\\u00e4te m\\u00fcssen im selben WLAN sein.',
      qr_warning: '\\u26a0\\ufe0f Der Server l\\u00e4uft nur auf <b>127.0.0.1</b> \\u2014 das Smartphone kann ihn nicht erreichen.<br>Starte mit <code>AGENT_HOST=0.0.0.0 python agent.py</code> f\\u00fcr Netzwerkzugriff.',
      setup_title:    '\\u2699 Agent einrichten',
      setup_desc:     'Gib deinem Agenten einen Namen und beschreibe dein Ziel. Diese Infos werden dauerhaft gespeichert.',
      label_agent_name: 'AGENT-NAME',
      name_placeholder: 'z.B. ARIA',
      label_goal:       'DEIN ZIEL / KONTEXT',
      goal_placeholder: 'z.B. Ich bin Entwickler und arbeite an Python-Projekten.',
      btn_save: 'Speichern',
      role_user:            'Du',
      upload_img_prefix:    '[Bild hochgeladen: ',
      upload_file_prefix:   'Datei hochgeladen: ',
      upload_analyze:       ' \\u2014 bitte analysiere sie.',
      analyze_img:          'Analysiere dieses Bild: ',
      upload_error:         '\\u26a0\\ufe0f Upload fehlgeschlagen: ',
      reminder_prefix:      '📅 Erinnerung: ',
      connection_error:     '\\u26a0\\ufe0f Verbindungsfehler: ',
      save_error:           '\\u26a0\\ufe0f Fehler beim Speichern: ',
      saved_feedback:       '\\u2713 Gespeichert',
      new_memory_prompt:    'Neue Erinnerung eingeben:',
      file_truncated:       '\\n[... gek\\u00fcrzt]',
      file_prefix:          'Datei: ',
    }
  };

  function setLang(lang) {
    currentLang = lang;
    localStorage.setItem('lang', lang);
    // Update lang toggle button (shows the OTHER language you can switch to)
    document.getElementById('lang-btn').textContent = lang === 'en' ? 'DE' : 'EN';
    // textContent replacements
    document.querySelectorAll('[data-i18n]').forEach(el => {
      const key = el.dataset.i18n;
      if (STRINGS[lang][key] !== undefined) el.textContent = STRINGS[lang][key];
    });
    // title attribute replacements
    document.querySelectorAll('[data-i18n-title]').forEach(el => {
      const key = el.dataset.i18nTitle;
      if (STRINGS[lang][key] !== undefined) el.title = STRINGS[lang][key];
    });
    // placeholder attribute replacements
    document.querySelectorAll('[data-i18n-placeholder]').forEach(el => {
      const key = el.dataset.i18nPlaceholder;
      if (STRINGS[lang][key] !== undefined) el.placeholder = STRINGS[lang][key];
    });
    // innerHTML replacements (trusted static strings only)
    document.querySelectorAll('[data-i18n-html]').forEach(el => {
      const key = el.dataset.i18nHtml;
      if (STRINGS[lang][key] !== undefined) el.innerHTML = STRINGS[lang][key];
    });
  }

  function toggleLang() {
    setLang(currentLang === 'en' ? 'de' : 'en');
  }

  // ── Init ──────────────────────────────────────────────────────────────────
  async function init() {
    setLang(currentLang);  // apply saved/default language before anything renders
    const [status, cfg] = await Promise.all([
      fetch('/status').then(r => r.json()),
      fetch('/config').then(r => r.json()),
    ]);
    document.getElementById('model-name').textContent = status.model;
    document.getElementById('root-path').textContent  = status.root;

    if (cfg.agent_name) {
      agentName = cfg.agent_name;
      document.getElementById('agent-label').textContent = cfg.agent_name;
    } else {
      // First run — show setup modal
      openSetup();
    }
  }
  init();

  // ── Notifications polling (every 30 s) ───────────────────────────────────
  async function pollNotifications() {
    try {
      const data = await fetch('/notifications').then(r => r.json());
      data.notifications.forEach(n => {
        showToast(n.title, n.description || n.due_at);
        const body = n.description ? `${n.description}` : '';
        appendMsg('agent', `${STRINGS[currentLang].reminder_prefix}${n.title} (${n.due_at})${body ? '\\n' + body : ''}`, []);
      });
    } catch(_) {}
  }
  pollNotifications();
  setInterval(pollNotifications, 30_000);

  // ── Toast ──────────────────────────────────────────────────────────────────
  function showToast(title, body) {
    const c = document.getElementById('toast-container');
    const t = document.createElement('div');
    t.className = 'toast';
    t.innerHTML = `<div class="toast-title">📅 ${escHtml(title)}</div>${escHtml(body)}`;
    c.appendChild(t);
    setTimeout(() => t.remove(), 7000);
  }

  // ── QR Modal ──────────────────────────────────────────────────────────────
  async function openQr() {
    const d = await fetch('/qrcode').then(r => r.json());
    document.getElementById('qr-img').src = 'data:image/png;base64,' + d.qr;
    document.getElementById('qr-url').textContent = d.url;
    document.getElementById('qr-warning').style.display = d.reachable ? 'none' : 'block';
    document.getElementById('qr-modal').classList.add('open');
  }

  // ── Directory Modal ───────────────────────────────────────────────────────
  function openDirModal() {
    const cur = document.getElementById('root-path').textContent;
    document.getElementById('dir-input').value = cur === '–' ? '' : cur;
    document.getElementById('dir-error').style.display = 'none';
    document.getElementById('dir-modal').classList.add('open');
    setTimeout(() => document.getElementById('dir-input').select(), 100);
  }

  async function saveDir() {
    const path = document.getElementById('dir-input').value.trim();
    if (!path) return;
    const errEl = document.getElementById('dir-error');
    const res = await fetch('/set-root', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ path }),
    });
    const data = await res.json();
    if (data.error) {
      errEl.textContent = data.error;
      errEl.style.display = 'block';
      return;
    }
    document.getElementById('root-path').textContent = data.root;
    closeModal('dir-modal');
    // Clear chat since history was reset
    chat.innerHTML = '';
    chat.appendChild(empty);
    empty.style.display = 'flex';
  }

  // ── Setup Modal ───────────────────────────────────────────────────────────
  function openSetup() {
    fetch('/config').then(r => r.json()).then(cfg => {
      document.getElementById('setup-name').value = cfg.agent_name || '';
      document.getElementById('setup-goal').value = cfg.user_goal  || '';
    });
    document.getElementById('setup-modal').classList.add('open');
  }

  async function saveSetup() {
    const name = document.getElementById('setup-name').value.trim();
    const goal = document.getElementById('setup-goal').value.trim();
    if (!name) { document.getElementById('setup-name').focus(); return; }
    await fetch('/setup', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ agent_name: name, user_goal: goal }),
    });
    agentName = name;
    document.getElementById('agent-label').textContent = name;
    closeModal('setup-modal');
  }

  function closeModal(id) {
    document.getElementById(id).classList.remove('open');
  }

  // Close modal on overlay click
  document.querySelectorAll('.modal-overlay').forEach(el => {
    el.addEventListener('click', e => { if (e.target === el) el.classList.remove('open'); });
  });

  // ── Chat helpers ──────────────────────────────────────────────────────────
  function autoResize(el) {
    el.style.height = 'auto';
    el.style.height = Math.min(el.scrollHeight, 160) + 'px';
  }

  function handleKey(e) {
    if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); sendMessage(); }
  }

  function sendHint(text) {
    input.value = text;
    autoResize(input);
    sendMessage();
  }

  // ── File upload / drag-drop ───────────────────────────────────────────────
  const TEXT_EXTS = new Set(['txt','md','py','js','ts','html','css','json','csv','xml',
                              'yaml','yml','toml','sh','bat','log','ini','cfg','sql']);

  async function handleFileSelect(e) {
    const file = e.target.files[0];
    e.target.value = '';
    if (!file) return;
    await processFile(file);
  }

  async function processFile(file) {
    const ext = file.name.split('.').pop().toLowerCase();
    if (TEXT_EXTS.has(ext)) {
      // Read client-side, paste into textarea
      const text = await file.text();
      const snippet = text.length > 4000 ? text.slice(0, 4000) + STRINGS[currentLang].file_truncated : text;
      input.value = (input.value ? input.value + '\\n\\n' : '')
        + STRINGS[currentLang].file_prefix + file.name + '\\n```' + ext + '\\n' + snippet + '\\n```';
      autoResize(input);
      input.focus();
    } else {
      // Upload to server
      const formData = new FormData();
      formData.append('file', file);
      try {
        const res  = await fetch('/upload', { method: 'POST', body: formData });
        const data = await res.json();
        if (file.type.startsWith('image/')) {
          appendMsg('user', STRINGS[currentLang].upload_img_prefix + file.name + ']', []);
          input.value = (input.value ? input.value + '\\n' : '')
            + STRINGS[currentLang].analyze_img + data.path;
        } else {
          input.value = (input.value ? input.value + '\\n' : '')
            + STRINGS[currentLang].upload_file_prefix + data.path + STRINGS[currentLang].upload_analyze;
        }
        autoResize(input);
        input.focus();
      } catch(err) {
        appendMsg('agent', STRINGS[currentLang].upload_error + err.message, []);
      }
    }
  }

  // Drag-and-drop onto chat area
  chat.addEventListener('dragover', e => { e.preventDefault(); chat.classList.add('drag-over'); });
  chat.addEventListener('dragleave', () => chat.classList.remove('drag-over'));
  chat.addEventListener('drop', e => {
    e.preventDefault();
    chat.classList.remove('drag-over');
    const file = e.dataTransfer.files[0];
    if (file) processFile(file);
  });

  function escHtml(s) {
    return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;').replace(/'/g,'&#39;');
  }

  const TOOL_ICONS = {
    list_directory:    '📁',
    read_file:         '📄',
    write_file:        '✍️',
    create_directory:  '📁',
    web_search:        '🔍',
    save_memory:       '🧠',
    list_memories:     '🧠',
    add_appointment:   '📅',
    list_appointments: '📅',
    create_docx:       '📝',
    create_xlsx:       '📊',
    analyse_data:      '📊',
    create_chart:      '📊',
    create_ascii_art:  '🎨',
  };

  function updateLoaderText(text) {
    const span = document.getElementById('loader-text');
    if (span) span.textContent = text;
  }

  function renderContent(text) {
    // Use sentinels to safely inject img tags without bypassing escaping.
    // Any LLM-injected '<img ...' text will still be escaped by escHtml.
    const imgs = [];
    function mkImg(src, alt) {
      // Validate src is a safe relative path, not a data: or javascript: URL
      const safeSrc = src.startsWith('/file?path=') ? src : '/file?path=' + encodeURIComponent(src);
      const tag = '<img class="chat-img" src="' + safeSrc
        + '" alt="' + escHtml(alt) + '" onclick="this.classList.toggle(\\'fullsize\\')">';
      const sentinel = '\\uE000img' + imgs.length + '\\uE001';
      imgs.push(tag);
      return sentinel;
    }
    // Handle [IMG:filename] markers from chart tool
    text = String(text).replace(/\\[IMG:([^\\]]+)\\]/g, (_, fname) =>
      mkImg('/file?path=' + encodeURIComponent(fname), 'chart'));
    // Handle markdown images: ![alt](path)
    text = text.replace(/!\\[([^\\]*)\\]\\(([^)]+)\\)/g, (_, alt, src) =>
      mkImg(src, alt));
    // Handle code blocks — split, escape, reassemble
    const parts = text.split(/(```[\\w]*\\n?[\\s\\S]*?```)/);
    const html = parts.map(part => {
      const m = part.match(/^```(\\w*)\\n?([\\s\\S]*?)```$/);
      if (m) {
        const lang = m[1] || 'code';
        const code = m[2];
        return '<div class="code-block"><div class="code-header">'
          + '<span class="code-lang">' + escHtml(lang) + '</span>'
          + '<button class="copy-btn" onclick="copyCode(this)">⎘ Copy</button>'
          + '</div><pre class="code-pre">' + escHtml(code) + '</pre></div>';
      }
      return '<span class="text-part">' + escHtml(part) + '</span>';
    }).join('');
    // Replace sentinels with safe img tags AFTER all escaping is done
    return html.replace(/\uE000img(\\d+)\uE001/g, (_, i) => imgs[+i] || '');
  }

  function copyCode(btn) {
    const pre = btn.closest('.code-block').querySelector('.code-pre');
    navigator.clipboard.writeText(pre.textContent).then(() => {
      btn.textContent = '✓ Copied';
      setTimeout(() => { btn.textContent = '⎘ Copy'; }, 1500);
    });
  }

  function appendMsg(role, text, toolTrace) {
    empty.style.display = 'none';
    const wrap = document.createElement('div');
    wrap.className = `msg ${role}`;

    const label = document.createElement('div');
    label.className = 'msg-label';
    label.textContent = role === 'user' ? STRINGS[currentLang].role_user : agentName;
    wrap.appendChild(label);

    if (toolTrace && toolTrace.length > 0) {
      const traceWrap = document.createElement('div');
      traceWrap.className = 'tool-trace';
      toolTrace.forEach(t => {
        const item = document.createElement('div');
        item.className = 'tool-item';
        const argsStr = Object.entries(t.args)
          .map(([k,v]) => `${k}=${JSON.stringify(v).slice(0,60)}`)
          .join(', ');
        item.innerHTML = `
          <div class="tool-header" onclick="this.parentElement.classList.toggle('open')">
            <span class="tool-icon">⚙</span>
            <span class="tool-name">${escHtml(t.tool)}</span>
            <span class="tool-args">${escHtml(argsStr)}</span>
            <span class="tool-chevron">▶</span>
          </div>
          <pre class="tool-result">${escHtml(t.result)}</pre>`;
        traceWrap.appendChild(item);
        // Auto-render chart images from create_chart results
        const imgMatch = t.result.match(/\\[IMG:([^\\]]+)\\]/);
        if (imgMatch) {
          const img = document.createElement('img');
          img.className = 'chat-img';
          img.src = '/file?path=' + encodeURIComponent(imgMatch[1]);
          img.alt = 'chart';
          img.onclick = () => img.classList.toggle('fullsize');
          traceWrap.appendChild(img);
        }
      });
      wrap.appendChild(traceWrap);
    }

    const bubble = document.createElement('div');
    bubble.className = 'bubble';
    bubble.innerHTML = renderContent(text);
    wrap.appendChild(bubble);

    chat.appendChild(wrap);
    chat.scrollTop = chat.scrollHeight;
  }

  function showLoader() {
    const el = document.createElement('div');
    el.className = 'msg agent';
    el.id = 'loader';
    el.innerHTML = `
      <div class="msg-label">${escHtml(agentName)}</div>
      <div class="loader">
        <div class="loader-dots"><span></span><span></span><span></span></div>
        <span id="loader-text"></span>
      </div>`;
    el.querySelector('#loader-text').textContent = STRINGS[currentLang].loader_thinking;
    chat.appendChild(el);
    chat.scrollTop = chat.scrollHeight;
  }

  function removeLoader() { document.getElementById('loader')?.remove(); }

  async function sendMessage() {
    const text = input.value.trim();
    if (!text || busy) return;
    busy = true;
    sendBtn.disabled = true;
    input.value = '';
    input.style.height = 'auto';
    appendMsg('user', text);
    showLoader();
    try {
      const response = await fetch('/chat-stream', {
        method:  'POST',
        headers: {'Content-Type': 'application/json'},
        body:    JSON.stringify({ message: text }),
      });
      if (!response.ok) throw new Error('HTTP ' + response.status);
      const reader  = response.body.getReader();
      const decoder = new TextDecoder();
      let   buffer  = '';
      let   answered = false;
      while (true) {
        const { done, value } = await reader.read();
        if (done) break;
        buffer += decoder.decode(value, { stream: true });
        const lines = buffer.split('\\n');
        buffer = lines.pop(); // keep incomplete last line
        for (const line of lines) {
          if (!line.startsWith('data: ')) continue;
          let event;
          try { event = JSON.parse(line.slice(6)); } catch(_) { continue; }
          if (event.type === 'tool_start') {
            const icon = TOOL_ICONS[event.tool] || '⚙';
            updateLoaderText(icon + '\\u2009' + event.tool + '…');
          } else if (event.type === 'confirm_search') {
            document.getElementById('search-confirm-query').textContent = event.query;
            document.getElementById('search-confirm-modal').classList.add('open');
          } else if (event.type === 'answer') {
            removeLoader();
            appendMsg('agent', event.text, event.tool_trace);
            answered = true;
          } else if (event.type === 'error') {
            removeLoader();
            appendMsg('agent', '⚠️ ' + event.text, []);
            answered = true;
          }
        }
      }
      if (!answered) { removeLoader(); }
    } catch(e) {
      removeLoader();
      appendMsg('agent', STRINGS[currentLang].connection_error + e.message, []);
    }
    busy = false;
    sendBtn.disabled = false;
    input.focus();
  }

  async function resetChat() {
    await fetch('/reset', { method: 'POST' });
    chat.innerHTML = '';
    chat.appendChild(empty);
    empty.style.display = 'flex';
    input.focus();
  }

  input.focus();

  // ── Memorize button ───────────────────────────────────────────────────────
  async function triggerMemorize() {
    if (busy) return;
    busy = true;
    sendBtn.disabled = true;
    const memBtn = document.getElementById('memorize-btn');
    memBtn.disabled = true;
    showLoader();
    try {
      const res  = await fetch('/memorize', { method: 'POST' });
      const data = await res.json();
      removeLoader();
      appendMsg('agent', data.answer, data.tool_trace);
    } catch(e) {
      removeLoader();
      appendMsg('agent', STRINGS[currentLang].save_error + e.message, []);
    }
    busy = false;
    sendBtn.disabled = false;
    memBtn.disabled  = false;
    input.focus();
  }

  // ── Appointments Modal ─────────────────────────────────────────────────────
  async function openAppointments() {
    await loadAppointments();
    document.getElementById('appt-modal').classList.add('open');
  }

  async function loadAppointments() {
    const data = await fetch('/appointments').then(r => r.json());
    const list = document.getElementById('appt-list');
    list.innerHTML = '';
    if (!data.appointments.length) {
      list.innerHTML = '<p style="color:var(--muted);font-size:13px;text-align:center;padding:24px"></p>';
      list.querySelector('p').textContent = STRINGS[currentLang].appt_empty;
      return;
    }
    data.appointments.forEach(a => {
      const card = document.createElement('div');
      card.className = 'appt-card' + (a.past ? ' past' : '');
      card.innerHTML = `
        <span class="appt-icon">${a.past ? '✅' : '📅'}</span>
        <div class="appt-body">
          <div class="appt-title"></div>
          <div class="appt-date"></div>
          <div class="appt-desc"></div>
        </div>
        <button class="btn-delete" onclick="deleteAppointment(${a.id},this)">✕</button>`;
      card.querySelector('.appt-title').textContent = a.title;
      card.querySelector('.appt-date').textContent  = a.due_at;
      card.querySelector('.appt-desc').textContent  = a.description || '';
      list.appendChild(card);
    });
  }

  async function deleteAppointment(id, btn) {
    await fetch('/appointments/' + id, { method: 'DELETE' });
    btn.closest('.appt-card').remove();
    const list = document.getElementById('appt-list');
    if (!list.querySelector('.appt-card')) {
      list.innerHTML = '<p style="color:var(--muted);font-size:13px;text-align:center;padding:24px"></p>';
      list.querySelector('p').textContent = STRINGS[currentLang].appt_empty;
    }
  }

  // ── Memory Modal ──────────────────────────────────────────────────────────
  async function openMemory() {
    await loadMemories();
    document.getElementById('memory-modal').classList.add('open');
  }

  async function loadMemories() {
    const data = await fetch('/memories').then(r => r.json());
    const list = document.getElementById('memory-list');
    list.innerHTML = '';
    if (!data.memories.length) {
      list.innerHTML = '<p style="color:var(--muted);font-size:13px;text-align:center;padding:24px"></p>';
      list.querySelector('p').textContent = STRINGS[currentLang].mem_empty;
      return;
    }
    data.memories.forEach(m => {
      const card = document.createElement('div');
      card.className = 'memory-card';
      card.dataset.id = m.id;
      const imp = Math.round((m.importance || 0.5) * 100);
      card.innerHTML = `
        <div class="memory-card-row">
          <span class="category-badge"></span>
          <div class="importance-bar" title="Wichtigkeit ${imp}%">
            <div class="importance-fill" style="width:${imp}%"></div>
          </div>
          <input type="number" class="mem-importance" min="0" max="1" step="0.1" value="${m.importance || 0.5}">
          <button class="btn-delete" onclick="deleteMemory(${m.id},this)">✕</button>
        </div>
        <div class="memory-card-row">
          <span class="mem-cat-label" style="font-family:var(--mono);font-size:10px;color:var(--muted);flex-shrink:0"></span>
          <input type="text" class="mem-category">
        </div>
        <textarea class="mem-content" rows="2"></textarea>
        <div class="memory-card-row" style="justify-content:space-between">
          <span class="memory-meta"></span>
          <button class="btn-save-mem" onclick="saveMemory(${m.id},this)"></button>
        </div>`;
      // Set values via .value / .textContent to avoid XSS
      card.querySelector('.category-badge').textContent  = m.category;
      card.querySelector('.mem-cat-label').textContent   = STRINGS[currentLang].label_category;
      card.querySelector('.btn-save-mem').textContent    = STRINGS[currentLang].btn_save_mem;
      card.querySelector('.mem-category').value           = m.category;
      card.querySelector('.mem-content').value            = m.content;
      card.querySelector('.memory-meta').textContent      = m.created_at;
      // Update importance bar when input changes
      const impInput = card.querySelector('.mem-importance');
      const impFill  = card.querySelector('.importance-fill');
      impInput.addEventListener('input', () => {
        impFill.style.width = Math.round(parseFloat(impInput.value || 0) * 100) + '%';
      });
      list.appendChild(card);
    });
  }

  async function saveMemory(id, btn) {
    const card       = btn.closest('.memory-card');
    const content    = card.querySelector('.mem-content').value.trim();
    const category   = card.querySelector('.mem-category').value.trim() || 'general';
    const importance = parseFloat(card.querySelector('.mem-importance').value) || 0.5;
    if (!content) return;
    await fetch('/memories/' + id, {
      method: 'PUT',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ content, category, importance }),
    });
    const orig = btn.textContent;
    btn.textContent = STRINGS[currentLang].saved_feedback;
    setTimeout(() => { btn.textContent = orig; }, 1500);
  }

  async function deleteMemory(id, btn) {
    await fetch('/memories/' + id, { method: 'DELETE' });
    btn.closest('.memory-card').remove();
    const list = document.getElementById('memory-list');
    if (!list.querySelector('.memory-card')) {
      list.innerHTML = '<p style="color:var(--muted);font-size:13px;text-align:center;padding:24px"></p>';
      list.querySelector('p').textContent = STRINGS[currentLang].mem_empty;
    }
  }

  async function addNewMemory() {
    const content = prompt(STRINGS[currentLang].new_memory_prompt);
    if (!content || !content.trim()) return;
    await fetch('/memories', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({ content: content.trim(), category: 'general', importance: 0.5 }),
    });
    await loadMemories();
  }

  // ── Search confirmation ───────────────────────────────────────────────────
  async function respondSearch(approved) {
    document.getElementById('search-confirm-modal').classList.remove('open');
    await fetch('/search-confirm', {
      method:  'POST',
      headers: {'Content-Type': 'application/json'},
      body:    JSON.stringify({ approved }),
    });
  }

  // ── History Modal ─────────────────────────────────────────────────────────
  async function openHistory() {
    await loadHistory();
    document.getElementById('history-modal').classList.add('open');
  }

  async function loadHistory() {
    const data = await fetch('/chat-log?limit=200').then(r => r.json());
    const list = document.getElementById('history-list');
    list.innerHTML = '';
    if (!data.entries.length) {
      list.innerHTML = '<p style="color:var(--muted);font-size:13px;text-align:center;padding:24px"></p>';
      list.querySelector('p').textContent = STRINGS[currentLang].history_empty;
      return;
    }
    data.entries.forEach(e => {
      const entry = document.createElement('div');
      entry.className = 'history-entry ' + (e.role === 'user' ? 'user-entry' : 'agent-entry');

      const roleRow = document.createElement('div');
      roleRow.className = 'history-role';

      const roleLabel = document.createElement('span');
      roleLabel.className = 'role-label';
      roleLabel.textContent = e.role === 'user' ? STRINGS[currentLang].role_user : agentName;

      const ts = document.createElement('span');
      ts.textContent = e.created_at;

      roleRow.appendChild(roleLabel);
      roleRow.appendChild(ts);

      const content = document.createElement('div');
      content.className = 'history-content';
      content.textContent = e.content;  // textContent — no XSS risk

      entry.appendChild(roleRow);
      entry.appendChild(content);
      list.appendChild(entry);
    });
  }

  async function clearHistory() {
    if (!confirm(STRINGS[currentLang].confirm_clear_history)) return;
    await fetch('/chat-log', { method: 'DELETE' });
    await loadHistory();
  }
</script>
</body>
</html>
"""

# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import uvicorn
    host = os.environ.get("AGENT_HOST", "127.0.0.1")
    print(f"🤖  Modell  : {MODEL}")
    print(f"📁  Root    : {ROOT}")
    print(f"🌐  URL     : http://{host}:8000")
    if host == "0.0.0.0":
        print(f"⚠️  LAN-Zugriff aktiv — keine Authentifizierung! LAN-IP: {_local_ip()}")
    print()
    uvicorn.run(app, host=host, port=8000)
